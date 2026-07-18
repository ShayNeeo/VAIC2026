from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Admin\Desktop\hakathon")
WORK = ROOT / "outputs" / "proposal_rewrite"
SOURCE = WORK / "source.docx"
OUTPUT = WORK / "SHB_Corporate_Sales_MVP_Data_Blueprint_V3_Proposal.docx"
ARCH_IMAGE = WORK / "context_aware_workflow.png"

BLUE = "17466B"
BLUE_2 = "2E74B5"
TEAL = "128C8C"
TEAL_LIGHT = "E8F5F4"
BLUE_LIGHT = "EAF2F8"
AMBER = "D98E04"
AMBER_LIGHT = "FFF4D6"
RED = "C83E3E"
RED_LIGHT = "FCEAEA"
GREEN = "2E8B57"
GREEN_LIGHT = "EAF6EF"
GRAY = "6B7280"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "1F2937"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_width(table, width_dxa: int = 9360) -> None:
    # Full-width tables are anchored to the text margin. Center alignment plus
    # a non-zero tblInd can make Word offset cell text on even PDF pages.
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: Sequence[int]) -> None:
    """Keep tblGrid, tblW and every tcW consistent for Word rendering."""
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_paragraph_shading(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE_2, 18, 10),
        ("Heading 2", 13, BLUE_2, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Caption" in styles:
        cap = styles["Caption"]
    else:
        cap = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = "Calibri"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(GRAY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(BLACK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.05


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("SHB CORPORATE SALES EXPERT WORKSPACE  •  MVP DATA BLUEPRINT V3")
        set_run_font(r, size=8.5, color=BLUE, bold=True)
        p.paragraph_format.space_after = Pt(0)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Tài liệu đề xuất – dữ liệu và chính sách minh họa cần được chủ sở hữu xác nhận  |  ")
        set_run_font(r, size=8, color=GRAY)
        add_page_field(p)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(155)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("SHB CORPORATE SALES\nEXPERT WORKSPACE")
    set_run_font(r, size=29, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("MVP bán hàng doanh nghiệp và Data Blueprint\ncho hệ thống Context-Aware Expert Workspace")
    set_run_font(r, size=15, color=TEAL, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Hiểu bối cảnh → hiểu đúng ý định → tìm giải pháp → kiểm tra điều kiện\n→ chuẩn bị case/task → RM phê duyệt → thực thi có kiểm soát")
    set_run_font(r, size=11.5, color=BLACK, bold=True)
    set_paragraph_shading(p, BLUE_LIGHT, BLUE_2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("Bản đề xuất V3 • 17/07/2026\nTrọng tâm: vertical slice có thể demo ngày 18/07/2026\nTác giả: Đào Quang Thắng")
    set_run_font(r, size=10, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("MVP/Hackathon Proposal – Không phải tài liệu chính sách chính thức của SHB")
    set_run_font(r, size=9, color=RED, bold=True)
    doc.add_page_break()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, *, bold_lead: str | None = None, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    num_id = None
    if numbered:
        # Create a distinct numbering instance so every numbered list starts at 1.
        numbering = doc.part.numbering_part.element
        base_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
        base_num = next(
            node for node in numbering.findall(qn("w:num"))
            if int(node.get(qn("w:numId"))) == int(base_num_id)
        )
        abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
        num_id = max(
            (int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))),
            default=0,
        ) + 1
        new_num = OxmlElement("w:num")
        new_num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_num_id))
        new_num.append(abstract_ref)
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        new_num.append(level_override)
        numbering.append(new_num)

    for item in items:
        p = doc.add_paragraph(item, style=style)
        if num_id is not None:
            num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
            num_pr.get_or_add_ilvl().val = 0
            num_pr.get_or_add_numId().val = num_id


def add_callout(doc: Document, title: str, body: str, *, fill=BLUE_LIGHT, border=BLUE_2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title + "\n")
    set_run_font(r, size=11, color=border, bold=True)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=BLACK)
    set_paragraph_shading(p, fill, border)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p.add_run(text)
    set_paragraph_shading(p, LIGHT, BLUE_2)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int] | None = None,
              font_size: float = 9.2, header_fill: str = LIGHT,
              keep_table_together: bool = False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(header)
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, size=font_size, color=BLUE, bold=True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    set_run_font(r, size=font_size, color=BLACK)
    if widths:
        set_table_grid(table, widths)
        for row in table.rows:
            for idx, w in enumerate(widths):
                set_cell_width(row.cells[idx], w)
    if keep_table_together:
        # Chaining each row to the next keeps a short table on one page. This
        # avoids a Microsoft Word PDF-export defect that can offset cell text
        # on continuation pages.
        for row in table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    doc.add_paragraph()


def copy_source_table(doc: Document, source: Document, index: int, caption: str, widths: Sequence[int] | None = None,
                      font_size: float = 8.8, caption_before: bool = False,
                      keep_table_together: bool = False) -> None:
    src = source.tables[index]
    headers = [c.text.strip().replace("\n", " / ") for c in src.rows[0].cells]
    rows = [[c.text.strip().replace("\n", " / ") for c in row.cells] for row in src.rows[1:]]
    if len(src.rows) == 1 and len(src.columns) == 1:
        add_callout(doc, caption, headers[0], fill=LIGHT, border=TEAL)
        return
    if caption_before:
        p = doc.add_paragraph(caption, style="Caption")
        p.paragraph_format.keep_with_next = True
    add_table(
        doc, headers, rows, widths=widths, font_size=font_size,
        keep_table_together=keep_table_together,
    )
    if not caption_before:
        p = doc.add_paragraph(caption, style="Caption")
        p.paragraph_format.keep_with_next = False


def create_architecture_image(path: Path) -> None:
    width, height = 2400, 1500
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    regular_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    font_title = ImageFont.truetype(bold_path, 54)
    font_sub = ImageFont.truetype(regular_path, 28)
    font_box = ImageFont.truetype(bold_path, 29)
    font_small = ImageFont.truetype(regular_path, 23)

    def color(hex_value: str):
        return "#" + hex_value

    def centered_text(box, text, font, fill, line_gap=8):
        x1, y1, x2, y2 = box
        lines = text.split("\n")
        metrics = [draw.textbbox((0, 0), line, font=font) for line in lines]
        heights = [b[3] - b[1] for b in metrics]
        total = sum(heights) + line_gap * (len(lines) - 1)
        y = y1 + (y2 - y1 - total) / 2
        for line, b, h in zip(lines, metrics, heights):
            tw = b[2] - b[0]
            draw.text((x1 + (x2 - x1 - tw) / 2, y), line, font=font, fill=fill)
            y += h + line_gap

    def box(x, y, w, h, title, subtitle="", outline=BLUE_2, fill=WHITE):
        rect = (x, y, x + w, y + h)
        draw.rounded_rectangle(rect, radius=20, fill=color(fill), outline=color(outline), width=5)
        if subtitle:
            centered_text((x + 10, y + 14, x + w - 10, y + h * 0.58), title, font_box, color(BLUE))
            centered_text((x + 15, y + h * 0.55, x + w - 15, y + h - 8), subtitle, font_small, color(GRAY))
        else:
            centered_text(rect, title, font_box, color(BLUE))
        return rect

    def arrow(start, end, fill=BLUE_2, width_px=6):
        draw.line([start, end], fill=color(fill), width=width_px)
        ex, ey = end
        sx, sy = start
        dx, dy = ex - sx, ey - sy
        length = max((dx * dx + dy * dy) ** 0.5, 1)
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        size = 20
        p1 = (ex, ey)
        p2 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
        p3 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
        draw.polygon([p1, p2, p3], fill=color(fill))

    centered_text((80, 30, width - 80, 110), "SHB CORPORATE EXPERT WORKSPACE — WORKFLOW V2", font_title, color(BLUE))
    centered_text((80, 105, width - 80, 155), "Context-first, intent-first, evidence-grounded và human-in-the-loop", font_sub, color(TEAL))

    a = box(105, 220, 400, 130, "RM WORKSPACE", "Yêu cầu ngắn + màn hình/case đang mở", TEAL, TEAL_LIGHT)
    b = box(610, 220, 470, 130, "CONTEXT ENGINE", "IAM • workspace • CRM • DMS • history", BLUE_2, BLUE_LIGHT)
    c = box(1190, 220, 470, 130, "INTENT RESOLVER", "Multi-intent • slots • provenance • confidence", BLUE_2, BLUE_LIGHT)
    d = box(1770, 220, 500, 130, "WORKFLOW ORCHESTRATOR", "Router • DAG • state • partial resume", BLUE, LIGHT)
    arrow((505, 285), (610, 285), TEAL)
    arrow((1080, 285), (1190, 285), BLUE_2)
    arrow((1660, 285), (1770, 285), BLUE_2)

    product = box(160, 540, 580, 170, "PRODUCT KNOWLEDGE", "Ingestion • hybrid RAG • metadata filter\nproduct matching + citations", BLUE_2, BLUE_LIGHT)
    legal = box(910, 540, 580, 170, "ELIGIBILITY & LEGAL", "Deterministic rules • Legal RAG\nKYC/UBO • missing information", AMBER, AMBER_LIGHT)
    ops = box(1660, 540, 580, 170, "OPERATIONS", "Checklist • brief • email/case/task draft\ndedup • artifact reuse • SLA", TEAL, TEAL_LIGHT)
    arrow((2020, 350), (450, 540), BLUE_2)
    arrow((2020, 350), (1200, 540), BLUE_2)
    arrow((2020, 350), (1950, 540), BLUE_2)

    evidence = box(470, 860, 620, 150, "EVIDENCE & SAFETY GATE", "Nguồn/phiên bản • claim support • ACL • PII\nchặn hallucination và tool vượt quyền", GREEN, GREEN_LIGHT)
    approval = box(1310, 860, 620, 150, "RM REVIEW & APPROVAL", "Hiển thị context, payload diff và rủi ro\ntoken gắn payload hash, expiry, one-time use", RED, RED_LIGHT)
    arrow((450, 710), (680, 860), GREEN)
    arrow((1200, 710), (780, 860), GREEN)
    arrow((1950, 710), (1780, 860), GREEN)
    arrow((1090, 935), (1310, 935), GREEN)

    execute = box(885, 1160, 630, 150, "CONTROLLED EXECUTOR", "CRM case/task/email chỉ chạy sau approval\nidempotency • reconciliation • immutable audit", BLUE, LIGHT)
    arrow((1620, 1010), (1200, 1160), RED)

    centered_text((120, 1370, width - 120, 1440), "Shared Case State • PostgreSQL/Vector DB/Cache • Audit • Trace • Metrics • Versioned contracts", font_sub, color(GRAY))
    img.save(path, quality=95)


def add_contents(doc: Document) -> None:
    add_h(doc, "TÓM TẮT ĐIỀU HÀNH", 1)
    add_p(doc, "SHB Corporate Expert Workspace là một không gian làm việc nội bộ đóng vai trò như “đội chuyên gia số” đứng sau mỗi RM. Hệ thống không chỉ trả lời câu hỏi, mà hiểu RM đang làm gì, đang phục vụ doanh nghiệp nào, case đang ở bước nào và dữ liệu nào đã tồn tại; từ đó hệ thống có thể tìm giải pháp, kiểm tra điều kiện, chuẩn bị checklist/case/task và soạn phản hồi mà không buộc nhân viên nhập lại cùng một thông tin qua nhiều vòng.")
    add_p(doc, "Bản V3 thu hẹp câu chuyện vào một vertical slice bán hàng doanh nghiệp có thể chạy và trình diễn ngày 18/07/2026: một doanh nghiệp 500 nhân sự, nhiều nhà cung cấp, dòng tiền phân tán và cần bổ sung vốn ngắn hạn. Hệ thống phải hiểu đồng thời bốn nhu cầu — chi lương, thanh toán nhà cung cấp, quản lý dòng tiền và vốn lưu động — rồi tạo một bundle giải pháp. Nhánh giao dịch được tiếp tục chuẩn bị; riêng nhánh tín dụng dừng ở pending_information nếu thiếu UBO hoặc báo cáo tài chính.")
    add_p(doc, "Product, Legal/Compliance, Operations, Evidence và Approval vẫn là các vai trò chuyên môn rõ ràng nhưng cùng đọc một case state có schema. Planner chỉ lập dependency và điều phối. Những quyết định rủi ro cao thuộc về rule deterministic và con người. Mọi thao tác tạo dữ liệu thật hoặc gửi ra ngoài phải qua evidence validation, RM review và approval gắn với đúng payload.")
    add_callout(doc, "Thông điệp cốt lõi", "RM không thiếu dữ liệu; RM cần một cơ chế phối hợp tri thức liên phòng ban để biến dữ liệu thành quyết định và hành động có kiểm soát. Hệ thống tối ưu phải hiểu context trước khi đọc câu lệnh, tái sử dụng trước khi tạo mới và chỉ hỏi khi thiếu dữ liệu thực sự làm thay đổi quyết định.", fill=TEAL_LIGHT, border=TEAL)
    add_callout(doc, "Scope freeze cho MVP ngày mai", "Không xây PDF ingestion, vector database persistent, CRM thật, email thật hoặc autonomous multi-agent trong sprint này. Dùng catalog/policy/SOP synthetic có version, deterministic rules, FastAPI và UI hiện có. Ưu tiên một hành trình hoàn chỉnh, dễ hiểu và có bằng chứng hơn số lượng tính năng.", fill=AMBER_LIGHT, border=AMBER)

    add_h(doc, "Mục tiêu kinh doanh", 2)
    add_bullets(doc, [
        "Rút ngắn thời gian từ lúc tiếp nhận nhu cầu đến khi có phương án tư vấn và checklist hành động.",
        "Giảm trao đổi lặp giữa RM với Product, Legal/Compliance và Operations cho các câu hỏi chuẩn hóa.",
        "Giảm câu hỏi làm rõ không cần thiết bằng cách dùng context đã có trong workspace, CRM, DMS, task và lịch sử case.",
        "Đảm bảo sản phẩm, điều kiện và phản hồi quan trọng đều có nguồn, phiên bản và trạng thái hiệu lực.",
        "Ngăn tạo trùng case/task, ngăn gửi sai nội dung và giữ đầy đủ audit trail.",
    ])
    add_h(doc, "Hai chỉ tiêu không được thỏa hiệp", 2)
    add_table(doc, ["Chỉ tiêu", "MVP", "Pilot"], [
        ["Unnecessary clarification rate", "< 10%", "< 5%"],
        ["Unsafe external action rate", "0%", "0%"],
        ["Duplicate task/action rate", "0%", "0%"],
        ["Important claims with valid evidence", "100%", "100%"],
    ], widths=[3900, 2730, 2730], font_size=9.5)

    add_h(doc, "Phạm vi và giả định", 2)
    add_bullets(doc, [
        "Người dùng trực tiếp là RM hoặc nhân viên kinh doanh khách hàng doanh nghiệp; khách hàng doanh nghiệp không trực tiếp tương tác với agent trong MVP.",
        "Dữ liệu sản phẩm, chính sách và quy trình thuộc nội bộ; MVP/hackathon dùng dữ liệu synthetic, không được coi là chính sách thật của SHB.",
        "Hệ thống không tự phê duyệt tín dụng, không tự chấp thuận khách hàng và không thay thế người có thẩm quyền.",
        "CRM/DMS/task/email được mô phỏng bằng adapter; tích hợp thật chỉ được thực hiện sau khi chốt API, IAM, retention và approval matrix.",
        "Tài liệu tải lên là dữ liệu không tin cậy; nội dung tài liệu không được phép thay đổi system policy hoặc quyền gọi tool.",
    ])

    add_h(doc, "CẤU TRÚC TÀI LIỆU", 1)
    add_table(doc, ["Phần", "Nội dung"], [
        ["1–3", "Bối cảnh, câu chuyện người dùng và định vị sản phẩm"],
        ["4–7", "Context, intent, confidence và chiến lược không hỏi lặp"],
        ["8–14", "Kiến trúc, chiến lược dữ liệu, Product RAG, Eligibility/Legal, Operations và Safety"],
        ["15–19", "API/UI, storage, observability, evaluation, lộ trình build và hiện trạng"],
        ["20–21", "Rủi ro, quyết định mở và kết luận"],
        ["Phụ lục", "Tool contract, state, backlog agent và catalog sản phẩm minh họa"],
    ], widths=[1550, 7810], font_size=9.5)
    doc.add_page_break()


def build_document() -> None:
    WORK.mkdir(parents=True, exist_ok=True)
    create_architecture_image(ARCH_IMAGE)
    source = Document(SOURCE)
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)
    doc.core_properties.title = "SHB Corporate Sales Expert Workspace – MVP Data Blueprint V3"
    doc.core_properties.subject = "Đề xuất hệ thống AI nội bộ hỗ trợ RM phục vụ khách hàng doanh nghiệp"
    doc.core_properties.author = "Đào Quang Thắng"
    doc.core_properties.keywords = "SHB, RM, Context, Intent, RAG, Multi-Agent, Human-in-the-loop"

    add_title(doc)
    add_contents(doc)

    add_h(doc, "1. BỐI CẢNH, ĐỐI TƯỢNG SỬ DỤNG VÀ PAIN POINT", 1)
    add_h(doc, "1.1. Ba lớp đối tượng cần phân biệt", 2)
    add_p(doc, "Giải pháp phục vụ một chuỗi giá trị gồm khách hàng doanh nghiệp, RM/nhân viên SHB và các đơn vị chuyên môn nội bộ. Việc phân biệt ba lớp này rất quan trọng: khách hàng nêu nhu cầu và cung cấp hồ sơ; RM chịu trách nhiệm quan hệ, phán đoán và phê duyệt nội dung; hệ thống AI chỉ hỗ trợ tổng hợp, đề xuất và chuẩn bị hành động.")
    copy_source_table(doc, source, 3, "Bảng 1. Đối tượng, vai trò và nhu cầu trong sản phẩm.", widths=[1900, 3400, 4060], font_size=8.8)

    add_h(doc, "1.2. Câu chuyện công việc hiện tại", 2)
    add_p(doc, "Một khách hàng doanh nghiệp hiếm khi chỉ hỏi về một sản phẩm. Khi doanh nghiệp nói “chúng tôi muốn trả lương, thu tiền từ đại lý, quản lý dòng tiền và có hạn mức khi thiếu hụt”, RM phải đồng thời hiểu mô hình vận hành, tra cứu nhiều dòng sản phẩm, kiểm tra hồ sơ pháp lý, xác định điều kiện tín dụng, lập danh sách hồ sơ và phối hợp các bộ phận để tạo task. Mỗi bước thường nằm ở một tài liệu, một hệ thống hoặc một đầu mối khác nhau.")
    add_bullets(doc, [
        "RM thu thập thông tin doanh nghiệp và diễn giải nhu cầu từ ngôn ngữ tự nhiên.",
        "RM tra cứu catalog, biểu phí, chính sách hoặc liên hệ Product để tạo bộ giải pháp.",
        "Legal/Compliance kiểm tra đăng ký doanh nghiệp, người đại diện, UBO/KYC, hiệu lực và ngoại lệ.",
        "Operations kiểm tra checklist, SOP, chủ sở hữu công việc, SLA và trạng thái xử lý.",
        "RM tổng hợp kết quả, hỏi bổ sung, tạo case/task, soạn phản hồi và theo dõi tiến độ.",
    ])
    add_p(doc, "Vấn đề không chỉ là thời gian tra cứu. Khi context bị thất lạc giữa các kênh, nhân viên phải nhắc lại khách hàng nào, case nào, đã có hồ sơ gì và đã hỏi gì. Công việc lặp lại phát sinh cả ở phía RM lẫn các bộ phận hỗ trợ; trong khi đó một chatbot hoặc RAG đơn chỉ trả lời được một câu hỏi nhưng không quản lý dependency, trạng thái và hành động tiếp theo.")
    copy_source_table(doc, source, 4, "Bảng 2. Hệ quả của quy trình phân mảnh.", widths=[2800, 6560], font_size=9)

    add_h(doc, "1.3. Root causes", 2)
    add_bullets(doc, [
        "Tri thức sản phẩm, pháp lý và vận hành phân tán, có phiên bản và thời hạn hiệu lực khác nhau.",
        "Không có shared case state thống nhất để ghi intent, facts, evidence, task, approval và lịch sử thay đổi.",
        "Ứng dụng chưa truyền đầy đủ workspace context sang trợ lý nên AI nhìn thấy câu lệnh nhưng không biết RM đang đứng ở đâu.",
        "LLM thường được kỳ vọng tự suy luận cả nghiệp vụ cứng, làm tăng rủi ro hallucination và khó audit.",
        "Thiếu cơ chế dedup, impact graph và partial resume nên cùng một kết quả bị tính lại hoặc cùng một task bị tạo lại.",
    ])

    add_h(doc, "2. ĐỊNH VỊ SẢN PHẨM VÀ JOB-TO-BE-DONE", 1)
    add_h(doc, "2.1. Định vị", 2)
    add_callout(doc, "SHB Corporate Expert Workspace", "Một workspace nội bộ giúp RM hiểu và xử lý yêu cầu khách hàng doanh nghiệp theo một case thống nhất. Hệ thống đóng vai trò đội chuyên gia số: Product tìm bộ giải pháp, Legal/Compliance kiểm tra điều kiện, Operations chuẩn bị hành động; Planner điều phối, Evidence Validator kiểm căn cứ và RM giữ quyền quyết định cuối.", fill=BLUE_LIGHT, border=BLUE)
    add_p(doc, "Định vị này khác chatbot hỏi–đáp. Sản phẩm phải tạo được một outcome nghiệp vụ có thể kiểm chứng: RM thấy hệ thống đang hiểu gì, sản phẩm nào được đề xuất, điều kiện nào đã đạt/chưa đạt, dữ liệu nào còn thiếu, task nào đã tồn tại và payload nào sẽ được tạo nếu phê duyệt.")

    add_h(doc, "2.2. Job-to-be-done", 2)
    add_p(doc, "Khi RM đang xử lý một doanh nghiệp, hệ thống phải tự nhận biết context hiện tại, hiểu yêu cầu ngắn hoặc mơ hồ, tái sử dụng dữ liệu và kết quả đã có, sau đó chuẩn bị bước tiếp theo mà không tạo công việc trùng. Công thức sản phẩm là:")
    add_code(doc, "Nắm bắt context nhân viên\n→ hiểu intent và tự điền dữ liệu có sẵn\n→ tìm sản phẩm có nguồn\n→ kiểm tra điều kiện bằng rule + evidence\n→ chuẩn bị/update checklist, case, task và phản hồi\n→ RM xem, sửa và phê duyệt\n→ executor thực thi đúng payload một lần")

    add_h(doc, "2.3. Phạm vi không làm", 2)
    add_bullets(doc, [
        "Không trực tiếp thay RM cam kết với khách hàng.",
        "Không tự quyết định hoặc phê duyệt tín dụng.",
        "Không suy đoán customer ID, product ID, người nhận, phí, hạn mức hoặc điều kiện khi không có nguồn.",
        "Không lưu suy luận nhạy cảm về tính cách hay hiệu suất nhân viên chỉ vì có behavioral context.",
        "Không để agent gọi trực tiếp CRM/email; mọi write đi qua Tool Gateway, Approval Service và Action Executor.",
        "Không bắt đầu bằng autonomous multi-agent chạy vòng lặp mở; workflow có trạng thái và điều kiện dừng là mặc định.",
    ])

    add_h(doc, "3. CÂU CHUYỆN END-TO-END: CÔNG TY ABC", 1)
    add_h(doc, "3.1. Tình huống", 2)
    add_callout(doc, "Case minh họa", "Công ty ABC là doanh nghiệp sản xuất có 500 nhân viên, nhiều nhà cung cấp và dòng tiền phân tán. Khách hàng muốn SHB tư vấn chi lương, thu/chi hộ, quản lý dòng tiền và vốn lưu động. Hồ sơ hiện có gồm đăng ký doanh nghiệp và CCCD người đại diện; thông tin UBO và báo cáo tài chính gần nhất chưa có.", fill=AMBER_LIGHT, border=AMBER)
    add_p(doc, "RM đang mở đúng customer/case trên Workspace và nhập: “Khách muốn chi lương, gom dòng tiền và có hạn mức khi thiếu hụt. Kiểm tra giúp tôi và soạn phản hồi hồ sơ còn thiếu.” Đây là một yêu cầu multi-intent: vừa tìm sản phẩm, vừa kiểm tra điều kiện, vừa chuẩn bị operations output.")

    add_h(doc, "3.2. Hệ thống phải xử lý như thế nào", 2)
    add_bullets(doc, [
        "Context Engine lấy employee ID, quyền, customer/case đang chọn, tài liệu hiện có và task đang mở; không hỏi lại “khách hàng nào?”.",
        "Intent Resolver tách Payroll, Cash Management, Working Capital, kiểm tra eligibility và soạn phản hồi; mỗi slot có source và confidence.",
        "Product RAG tìm sản phẩm trong catalog kiểm soát, lọc phiên bản còn hiệu lực và trả citation.",
        "Eligibility Engine chạy rule UBO/BCTC; Legal RAG chỉ cung cấp căn cứ và giải thích, không tự phê duyệt.",
        "Planner giữ các nhánh transaction services tiếp tục, chỉ chặn nhánh tín dụng ở pending_information.",
        "Operations hợp nhất checklist, phát hiện task UBO đã tồn tại hay chưa, cập nhật draft thay vì tạo trùng.",
        "Evidence Validator chặn claim “đủ điều kiện hoàn toàn”; RM thấy lý do và nguồn.",
        "RM duyệt payload. Executor tạo đúng case/task một lần với idempotency key và ghi audit.",
    ], numbered=True)
    copy_source_table(doc, source, 26, "Bảng 3. Luồng demo end-to-end trong đề xuất ban đầu, được giữ lại và đặt trong kiến trúc V2.", widths=[850, 2100, 6410], font_size=8.5)
    add_callout(doc, "“Wow moment” cần chứng minh", "Product tìm được bộ giải pháp nhưng Legal phát hiện thiếu UBO; Planner tự thay đổi kế hoạch, giữ nhánh an toàn tiếp tục và Operations chuẩn bị đúng action bổ sung. Đây là collaboration làm thay đổi workflow, không phải ba chatbot trả lời nối tiếp.", fill=GREEN_LIGHT, border=GREEN)

    add_h(doc, "3.3. Resume sau khi có UBO", 2)
    add_p(doc, "Khi RM hoặc khách hàng tải lên tài liệu UBO, hệ thống không chạy lại toàn bộ case. DMS phát event, workflow tính impact từ loại tài liệu và chỉ chạy lại Eligibility → Evidence → Operations. Intent và Product được giữ nguyên nếu input hash, catalog version và mục tiêu không đổi. Checklist/email/task cũ được cập nhật theo version; approval cũ bị vô hiệu nếu payload thay đổi.")

    add_h(doc, "4. CONTEXT-AWARE: HIỂU NHÂN VIÊN ĐANG LÀM GÌ", 1)
    add_h(doc, "4.1. Vì sao context phải đứng trước prompt", 2)
    add_p(doc, "Một câu như “Kiểm tra còn thiếu gì” không đủ nghĩa nếu tách khỏi màn hình và case. Nhưng nếu RM đang ở customer COMP-ABC, case Working Capital Review, tab Financial Documents và task Kiểm tra BCTC, intent gần như đã rõ. Hệ thống phải tải context trước khi gọi LLM; LLM nhận một snapshot đã chuẩn hóa và tối thiểu hóa, không nhận dump toàn bộ CRM.")

    add_h(doc, "4.2. Tám lớp context", 2)
    add_table(doc, ["Lớp", "Nội dung", "Nguồn", "Freshness mặc định"], [
        ["Employee", "employee_id, role, org unit", "SSO/HRIS", "Session / 24h"],
        ["Permission", "scopes, managed customers", "IAM", "5 phút"],
        ["Workspace", "screen, selected customer/case/task/product", "UI session", "Realtime"],
        ["Customer", "profile, segment, KYC, products", "CRM", "5 phút"],
        ["Workflow", "current node, open questions, task/artifact", "State DB", "Realtime"],
        ["Documents", "type, version, status, access", "DMS", "5 phút"],
        ["Conversation", "goal, confirmed facts, rejected assumptions", "State DB", "Session"],
        ["Preference", "language, brief/email format", "User settings", "30 ngày"],
    ], widths=[1500, 3430, 1950, 2480], font_size=8.6)
    copy_source_table(doc, source, 33, "Bảng 4. Dữ liệu nhân viên và nguồn đề xuất.", widths=[1900, 4300, 3160], font_size=8.6)

    add_h(doc, "4.3. Trình tự thu thập và ranh giới quyền", 2)
    add_code(doc, "Authenticated employee\n→ load IAM scope\n→ read workspace selection\n→ validate access to selected customer/case\n→ load CRM/task/document metadata in parallel\n→ load confirmed conversation facts\n→ normalize + timestamp + provenance\n→ detect conflicts and minimize context\n→ return ContextSnapshot")
    add_p(doc, "Nếu customer đang chọn không thuộc scope, hệ thống phải fail closed với lỗi truy cập; tuyệt đối không fallback sang customer gần nhất. Permission/IAM luôn thắng user input. Context đưa vào model chỉ gồm các field cần cho intent hiện tại; không gửi toàn bộ danh sách khách hàng, email, giấy tờ định danh hoặc nội dung case khác.")

    add_h(doc, "4.4. Quy tắc ưu tiên và xung đột", 2)
    add_bullets(doc, [
        "Giá trị user nêu rõ trong message hiện tại và hợp lệ.",
        "Workspace selection hiện tại.",
        "CRM/DMS/workflow còn fresh.",
        "Conversation fact đã được xác nhận.",
        "Cache còn TTL.",
        "LLM inference – chỉ cho field rủi ro thấp và không được ghi đè giá trị hệ thống/xác nhận.",
    ], numbered=True)
    add_p(doc, "Xung đột high-impact như customer, case, recipient hoặc product gắn external action phải được hiển thị và xác nhận trước write. Mọi field auto-fill phải có value, source, confidence, freshness và confirmed flag để UI giải thích được “AI lấy thông tin này từ đâu”.")

    add_h(doc, "5. INTENT UNDERSTANDING: HIỂU ĐÚNG VIỆC RM MUỐN LÀM", 1)
    add_h(doc, "5.1. Intent không chỉ là một label", 2)
    add_p(doc, "IntentResult phải biểu diễn job-to-be-done, sub-intents, target entities, action yêu cầu, constraints, success criteria, outputs, missing slots, ambiguity, evidence spans và field-level confidence. LLM chỉ làm semantic extraction; entity resolver, permission, ID normalization và workflow dependency là code/tool deterministic.")
    add_table(doc, ["Intent ID", "Ý nghĩa", "Slot chính", "Rủi ro"], [
        ["find_product", "Tìm giải pháp phù hợp", "customer/profile, objective", "Thấp"],
        ["compare_products", "So sánh ứng viên", "product candidates", "Thấp"],
        ["check_eligibility", "Kiểm tra điều kiện", "customer, product", "Trung bình/Cao"],
        ["check_missing_documents", "Kiểm tra hồ sơ thiếu", "case/customer, workflow/product", "Trung bình"],
        ["resume_case", "Tiếp tục sau cập nhật", "case, changed artifact", "Trung bình"],
        ["prepare_customer_response", "Soạn phản hồi", "case, purpose, recipient candidate", "Trung bình"],
        ["prepare_case_task", "Chuẩn bị case/task", "customer/case, task type", "Trung bình"],
        ["approve_actions", "Phê duyệt action", "case, frozen payload", "Cao"],
        ["status_lookup / out_of_scope", "Xem trạng thái / ngoài phạm vi", "case/task / none", "Thấp–biến đổi"],
    ], widths=[2200, 2940, 2780, 1440], font_size=8.4)

    add_h(doc, "5.2. Pipeline extraction", 2)
    add_code(doc, "Normalize Vietnamese text/abbreviations\n→ taxonomy + minimized ContextSnapshot\n→ LLM structured output\n→ JSON schema validation\n→ deterministic entity normalization\n→ merge slots from context/tools\n→ ambiguity/conflict calculation\n→ confidence and clarification policy")
    add_bullets(doc, [
        "Không tự tạo customer, product, amount, date hoặc urgency.",
        "Tách nhiều intent khi một câu chứa nhiều mục tiêu.",
        "Giữ nguyên số tiền, ngày, tên thực thể và evidence span từ message.",
        "Không biết thì để null/missing; product alias chỉ trở thành product ID khi catalog resolver chứng minh.",
        "Không lưu chain-of-thought; chỉ lưu rationale ngắn, evidence span và decision code phục vụ audit.",
    ])

    add_h(doc, "5.3. Ví dụ output đã resolve", 2)
    add_code(doc, '{\n  "primary_intent": "check_missing_documents",\n  "sub_intents": ["prepare_customer_response"],\n  "target_customer_id": {"value":"COMP-ABC","source":"workspace","confidence":1.0},\n  "active_case_id": {"value":"CASE-001","source":"workspace","confidence":1.0},\n  "required_outputs": ["missing_document_checklist","customer_email_draft"],\n  "unresolved_slots": [],\n  "recommended_action": "continue_workflow"\n}')

    add_h(doc, "6. SLOT AUTO-FILL, CONFIDENCE VÀ CHIẾN LƯỢC KHÔNG HỎI LẶP", 1)
    add_h(doc, "6.1. Resolution order", 2)
    add_code(doc, "user_explicit → workspace → workflow/case state → CRM/DMS\n→ conversation_confirmed → valid cache → deterministic derivation\n→ low-risk LLM inference → unresolved")
    add_p(doc, "Hệ thống không đặt mục tiêu tuyệt đối “không bao giờ hỏi”. Mục tiêu đúng là không hỏi lại dữ liệu có thể lấy được, không hỏi field chưa cần cho bước hiện tại và chỉ hỏi một câu có information gain cao nhất khi thiếu dữ liệu quyết định. Với external action, preview và explicit approval vẫn bắt buộc dù confidence cao.")
    copy_source_table(doc, source, 34, "Bảng 5. Trường có thể tự điền và trường không được suy đoán.", widths=[3000, 1800, 4560], font_size=8.8)

    add_h(doc, "6.2. Required-now và required-later", 2)
    add_p(doc, "Mỗi slot phải khai báo required_for_understanding, required_for_retrieval, required_for_eligibility và required_for_external_action. Ví dụ requested_amount có thể chưa cần để tìm product candidates nhưng bắt buộc ở bước tạo một số credit case. Workflow tiếp tục các bước an toàn, defer câu hỏi và chỉ dừng đúng node bị block.")

    add_h(doc, "6.3. Confidence policy", 2)
    add_table(doc, ["Nguồn", "Base confidence", "Ghi chú"], [
        ["Authenticated IAM/SSO", "1.00", "Quyền vẫn phải được kiểm tra theo scope"],
        ["Workspace selected ID", "1.00", "Có thể phát sinh conflict khi user switch"],
        ["Fresh CRM/DMS", "0.98", "Giảm điểm nếu stale"],
        ["User explicit current message", "0.95", "Không thể tự khai báo quyền"],
        ["Workflow state", "0.95", "Phải cùng case/version"],
        ["Conversation confirmed", "0.90", "Có evidence message"],
        ["Fresh cache / deterministic derivation", "0.85", "Key gồm version và scope"],
        ["LLM inference", "≤ 0.70", "Không dùng làm field quyết định cho write"],
    ], widths=[3300, 1700, 4360], font_size=8.7)
    copy_source_table(doc, source, 35, "Bảng 6. Ma trận confidence/risk/action.", widths=[2200, 2100, 5060], font_size=8.8)

    add_h(doc, "6.4. Clarification tối ưu", 2)
    add_bullets(doc, [
        "Liệt kê unresolved slots có decision impact.",
        "Xếp hạng theo information gain × risk × downstream blocking.",
        "Tự gọi read tool trước; không hỏi field đã có trong hệ thống nguồn.",
        "Hỏi tối đa một câu mỗi lượt, ưu tiên lựa chọn cụ thể khi có 2–3 hypothesis.",
        "Lưu câu trả lời thành confirmed fact có provenance.",
        "Khi user sửa context, tạo correction event, invalidate đúng descendants và resume từ node sớm nhất bị ảnh hưởng.",
    ], numbered=True)

    add_h(doc, "7. CÁC JOURNEY CHÍNH VÀ TIÊU CHÍ TRẢI NGHIỆM", 1)
    add_table(doc, ["Journey", "Input", "Hành vi bắt buộc"], [
        ["Kiểm tra hồ sơ case đang mở", "“Kiểm tra còn thiếu gì”", "Tự lấy customer/case/product; trả checklist/evidence; không hỏi lại; không tạo task trùng"],
        ["Nhu cầu đa sản phẩm", "Payroll + dòng tiền + hạn mức", "Tách multi-intent; block riêng credit nếu thiếu dữ liệu; giữ nhánh an toàn"],
        ["Resume sau upload", "UBO/BCTC mới", "Chạy lại Legal/Evidence/Ops; giữ intent/product; update artifact"],
        ["Sửa context", "RM đổi customer/product", "Hiển thị impact; invalidate descendants; không xóa audit"],
        ["External action", "RM bấm Approve", "Payload diff; token hash/expiry; RBAC/evidence/idempotency; execute một lần"],
    ], widths=[2050, 2450, 4860], font_size=8.6)

    add_h(doc, "8. KIẾN TRÚC TỔNG THỂ", 1)
    doc.add_picture(str(ARCH_IMAGE), width=Inches(6.45))
    p = doc.add_paragraph("Hình 1. Kiến trúc context-aware workflow V2: specialized agents/modules nằm trong một orchestration có state, evidence và approval.", style="Caption")
    add_h(doc, "8.1. Các lớp kiến trúc", 2)
    add_table(doc, ["Lớp", "Trách nhiệm"], [
        ["Experience", "RM Workspace, Context Header, Intent Preview, Evidence, Operations, Approval, Timeline"],
        ["API", "Context, case, document, workflow, approval, search; typed contract và auth principal"],
        ["Understanding", "Context Assembler, Intent Extractor, Slot Resolver, Confidence/Clarification"],
        ["Orchestration", "Complexity Router, Planner DAG, state machine, retry, impact-based resume"],
        ["Knowledge & Rules", "Product ingestion/RAG/matcher, Legal RAG, deterministic eligibility rules"],
        ["Operations & Safety", "Checklist, draft, dedup, artifact reuse, evidence, approval, executor"],
        ["Integration", "SSO/IAM, CRM, DMS, task, email, model gateway qua adapters"],
        ["Storage & Ops", "PostgreSQL, vector DB, cache, object store, audit, traces, metrics"],
    ], widths=[2100, 7260], font_size=9)

    add_h(doc, "8.2. Vì sao không dùng autonomous multi-agent làm mặc định", 2)
    add_p(doc, "Multi-agent có giá trị khi các vai trò có dữ liệu, tool, output contract và dependency khác nhau. Tuy nhiên các bước nghiệp vụ, trạng thái và điểm phê duyệt phải deterministic để retry, resume và audit. Vì vậy kiến trúc đề xuất giữ tên Product Agent, Legal Agent và Operations Agent ở tầng sản phẩm, nhưng runtime triển khai chúng như module/node typed; chỉ thêm planner reasoning ở case thực sự đa nhánh.")
    copy_source_table(doc, source, 6, "Bảng 7. Routing giữa yêu cầu đơn giản và phức tạp.", widths=[2100, 3100, 4160], font_size=8.8)
    copy_source_table(doc, source, 7, "Bảng 8. Chức năng, input và output của các khối trong proposal gốc.", widths=[1450, 2850, 2500, 2560], font_size=7.9)

    add_h(doc, "8.3. Shared contracts", 2)
    add_p(doc, "Mọi module giao tiếp qua typed state/command, không truyền dict tùy ý. ID chuẩn gồm case_id, trace_id, employee_id, customer_id, task_id và document_id. Mọi output có schema_version; mọi field suy luận có source/confidence/confirmed; mọi evidence có source/version/location; mọi external action có payload hash và idempotency key.")
    add_code(doc, "new → understanding → clarification_required → planned → in_analysis\n→ pending_information | pending_review | pending_approval\n→ executing → completed | rejected | failed")
    copy_source_table(doc, source, 24, "Bảng 9. Ý nghĩa và chuyển tiếp trạng thái nghiệp vụ.", widths=[1900, 3200, 4260], font_size=8.5)

    add_h(doc, "9. DATA STRATEGY VÀ MARKET DATA", 1)
    add_h(doc, "9.1. Câu hỏi phải trả lời trước khi xây AI", 2)
    add_p(doc, "Một solution context-aware/RAG chỉ tốt bằng dữ liệu mà nó có quyền sử dụng và có thể truy vết. Vì vậy đội dự án phải phân biệt bốn câu hỏi: dữ liệu có tồn tại không; có lấy được bằng kênh ổn định không; có quyền dùng cho mục đích AI không; và dữ liệu có đủ fresh/complete/provenance để ảnh hưởng quyết định không. “Có thể tìm thấy trên web” không đồng nghĩa “khả dụng cho production”.")
    add_bullets(doc, [
        "Dữ liệu nào là nguồn sự thật nội bộ bắt buộc và không thể mua ngoài?",
        "Dữ liệu official/open/commercial nào có thể xác minh hoặc làm giàu?",
        "Source nào được phép ảnh hưởng product matching, eligibility hoặc chỉ hiển thị tham khảo?",
        "Join key nào nối được external entity với customer_id mà không merge nhầm?",
        "Owner, legal basis/license, retention, data residency, update SLA và exit plan là gì?",
        "Pipeline nào biến raw source thành Gold artifact có version, ACL, quality và lineage?",
    ])

    add_h(doc, "9.2. Phân tầng nguồn và quyền quyết định", 2)
    add_table(doc, ["Tier", "Nguồn", "Ví dụ", "Vai trò được phép"], [
        ["A – Internal authoritative", "SHB sở hữu", "Product master, CRM, IAM, DMS, SOP, approved policy", "Nguồn chính nếu owner/version/freshness hợp lệ"],
        ["A – Official authoritative", "Cơ quan nhà nước/quốc tế", "Đăng ký doanh nghiệp, VBPL, official sanction lists", "Xác minh trong phạm vi/terms được phép"],
        ["B – Licensed curated", "Vendor theo hợp đồng", "Business/credit data, PEP/adverse media", "Enrichment/screening theo policy đã duyệt"],
        ["C – Open/public", "Open data", "Macro, LEI, open company data", "Discovery/benchmark; không tự pass eligibility"],
        ["D – Derived", "Do hệ thống tạo", "Entity match, summary, score, embedding", "Chỉ dùng kèm lineage/model/version/validation"],
        ["E – Synthetic/labeled", "Dev/eval", "Demo companies, golden cases", "Test/evaluation; không trộn production facts"],
    ], widths=[1500, 1800, 3200, 2860], font_size=7.8)
    add_callout(doc, "Hard veto", "Source thiếu owner, legal basis/license, purpose, access method hợp lệ hoặc provenance không được publish vào serving layer dù tổng điểm chất lượng cao.", fill=RED_LIGHT, border=RED)

    add_h(doc, "9.3. Bản đồ dữ liệu cần cho solution", 2)
    add_table(doc, ["Domain", "Dữ liệu cần", "Nguồn ưu tiên", "Thị trường/ngoài SHB", "Vai trò"], [
        ["Employee/workspace", "Identity, role, scope, screen/customer/case/task", "SSO/IAM/HRIS/UI", "Không thể mua", "Context + RBAC"],
        ["Customer/case", "Master, relationship, products, interactions, tasks", "CRM/DWH/task", "Vendor chỉ enrich", "Canonical customer_id"],
        ["Documents", "Legal/financial files, type, version, status", "DMS/upload", "Có OCR/parser", "Extracted facts + checklist"],
        ["Product", "Catalog, ID, segment, fees/limits, prerequisites", "Product master/docs", "Không có nguồn ngoài đáng tin cho SHB", "Product RAG/matcher"],
        ["Legal/SOP", "Policy, approval matrix, process, SLA", "Legal/Compliance/Ops", "Public law chỉ bổ sung", "Rules + evidence"],
        ["Business registry", "Name, registration ID, address, representative, status", "National registry", "Public/info services", "Entity verification"],
        ["Credit", "Credit history/report/score", "CIC + internal credit", "Controlled/licensed", "Read-only risk input"],
        ["KYC/AML", "Sanctions, PEP, adverse media/watchlists", "Compliance + official/vendor", "Official lists + commercial feeds", "Screening/review"],
        ["Market/industry", "Macro, sector, trade, benchmarks", "NSO/SBV + vendor", "Public/licensed", "Context/benchmark"],
        ["Eval/feedback", "Intent/evidence/outcome/corrections", "Synthetic + approved samples", "Không có dataset đủ sát SHB", "Regression and calibration"],
    ], widths=[1250, 2350, 1900, 2100, 1760], font_size=7.2)

    add_h(doc, "9.4. Market data landscape hiện có (khảo sát 17/07/2026)", 2)
    add_p(doc, "Các nguồn dưới đây phù hợp để shortlist/POC. Availability được ghi nhận từ trang chính thức hoặc trang sản phẩm tại thời điểm khảo sát; trước tích hợp phải xác minh API, quota, giá, license, data processing terms và coverage bằng mẫu thực tế.")
    add_table(doc, ["Nhóm", "Nguồn hiện có", "Dữ liệu/format", "Cách dùng và giới hạn"], [
        ["Đăng ký doanh nghiệp Việt Nam", "Cổng thông tin quốc gia về đăng ký doanh nghiệp", "Tra cứu public các trường cơ bản; có information services", "Verify entity; không production scrape nếu chưa có quyền/API"],
        ["Văn bản pháp luật", "CSDL quốc gia VBPL + SBV", "Web/PDF, thuộc tính/hiệu lực", "Legal ingestion; monitor version/effective date"],
        ["Credit information", "CIC", "Credit report/information/rating products qua kênh kiểm soát", "Không phải open data; chỉ adapter được cấp quyền"],
        ["Macro Việt Nam", "National Statistics Office – NSDP", "GDP, CPI, trade, FX, IIP, industry/labor; Excel/SDMX", "Benchmark/context; không quyết customer eligibility"],
        ["Vietnam corporate intelligence", "FiinGroup và provider tương đương", "Business reports/API, risk score, industry/trade research", "Commercial POC: coverage, lineage, SLA, license, calibration"],
        ["Global legal entity", "GLEIF", "LEI Level 1/2, mapped IDs; free API/full/delta", "Enrich entity có LEI; không thay registry Việt Nam"],
        ["Cross-border company data", "OpenCorporates/provider tương đương", "Company data có source; open/commercial versioned API", "Candidate resolution; verify primary source/license"],
        ["Official sanctions/debarment", "UN, OFAC, World Bank", "XML/HTML/PDF/download/search", "Input screening; cần matching/update/false-positive workflow"],
        ["Curated KYC/AML", "LSEG World-Check và vendor tương đương", "Sanctions, PEP/RCA, adverse media, watchlists", "Licensed enrichment; compliance review, PII and cross-border due diligence"],
    ], widths=[1500, 2200, 2750, 2910], font_size=7.1)
    add_p(doc, "Những nguồn bên ngoài trên không thay thế dữ liệu nội bộ. Product name/fee/limit/policy/SOP của SHB phải đến từ data owner SHB. CIC/KYC/vendor response không được lưu hoặc đưa vào prompt ngoài phạm vi được phép; việc vendor cung cấp một score không cho phép LLM tự động phê duyệt hoặc từ chối khách hàng.")

    add_h(doc, "9.5. Data fitness score và quy trình chọn source/vendor", 2)
    add_table(doc, ["Dimension", "Trọng số", "Câu hỏi kiểm tra"], [
        ["Legal/license fit", "20", "Có quyền ingest, cache, derive, embed và hiển thị citation?"],
        ["Availability/integration", "15", "API/file/event, auth, quota, sandbox, uptime?"],
        ["Accuracy/provenance", "15", "Primary source, evidence và correction process?"],
        ["Freshness", "15", "Update cadence phù hợp decision window?"],
        ["Coverage", "15", "Segment/SME/Việt Nam/historical depth?"],
        ["Joinability", "10", "Registration/tax ID, LEI hoặc stable identifier?"],
        ["Cost/latency", "5", "Cost/case, bulk/API pricing, P95?"],
        ["Operational fit", "5", "Monitoring, support, changelog và exit/export?"],
    ], widths=[2450, 1300, 5610], font_size=8.4)
    add_bullets(doc, [
        "Chọn 100–500 synthetic/de-identified entities đại diện segment và xác định ground truth mẫu.",
        "Đo coverage, field completeness, match precision/recall, stale rate, latency và cost.",
        "Red-team tên Việt Nam có dấu/không dấu, viết tắt, tên gần giống và thay đổi địa chỉ.",
        "Kiểm điều khoản cho caching, embeddings, derived scores, retention, subprocessors và deletion evidence.",
        "Chạy shadow mode; SME/Compliance ký Data Source Acceptance Record trước khi source ảnh hưởng quyết định.",
    ], numbered=True)

    add_h(doc, "9.6. Data Source Card và inventory bắt buộc", 2)
    add_p(doc, "Mỗi source có một Source Card versioned gồm source_id/domain/tier, business owner/data steward/technical owner, purpose và prohibited uses, decision role, legal basis/license/DPA, sensitivity, residency/retention, access/auth/quota/SLA, schema/format, identifiers/join keys, freshness/stale behavior, quality gates, ingestion lineage, consumers và lifecycle status. Đây là contract để AI coding không tự kết nối một nguồn chỉ vì thấy URL hoặc file.")

    add_h(doc, "9.7. Data preparation pipeline", 2)
    add_code(doc, "Discover/Register source\n→ owner + legal/license/privacy assessment\n→ acquire raw to quarantine + manifest/hash\n→ malware/type/schema validation\n→ parse/OCR/table extraction\n→ normalize encoding, units, dates, identifiers\n→ entity resolution to internal canonical IDs\n→ quality profiling + source reconciliation\n→ PII classification/minimization/masking + ACL\n→ version/effective/change detection\n→ publish Silver normalized data\n→ publish Gold product/rule/context/eval artifacts\n→ chunk/index or compile rules\n→ acceptance tests + owner sign-off\n→ Serving with lineage, trace and monitoring")
    add_table(doc, ["Layer", "Nội dung", "Quy tắc"], [
        ["Quarantine/Raw", "File/API payload bất biến + manifest/hash", "Không đưa model dùng trực tiếp"],
        ["Silver", "Parsed, typed, normalized, canonical IDs, quality flags", "Giữ source record và lineage"],
        ["Gold", "Approved product/policy/rules/context/eval artifacts", "Owner/version/effective/ACL bắt buộc"],
        ["Serving", "API tables, vector/sparse index, rule registry, feature views", "Chỉ artifact pass gate"],
        ["Audit", "Ingest report, changes, failures, source decisions", "Append-only và retention controlled"],
    ], widths=[2000, 4000, 3360], font_size=8.5)

    doc.add_page_break()
    add_h(doc, "9.8. Cách xử lý theo loại dữ liệu", 2)
    add_table(doc, ["Loại", "Chuẩn bị bắt buộc"], [
        ["PDF/Word policy", "SHA-256; parser/OCR; giữ heading/page/table; effective dates; structure-aware chunks; citation sample QA"],
        ["Excel/catalog", "Schema/unit/currency; product ID; effective rows; duplicate/conflict; row-level provenance"],
        ["CRM/API", "Adapter normalization; canonical ID; freshness; field-level provenance; CDC/event hoặc TTL"],
        ["KYC/vendor response", "Source record/reference; match features/score; review status; expiry; không log raw PII"],
        ["Conversation", "Message-level facts/corrections; PII redaction; retention; không lưu raw vô thời hạn"],
        ["Task/artifact", "Canonical dedup key; input/output hash; version; supersedes/reuse link"],
        ["Eval label", "Dataset version; expected IDs/outcomes/evidence; reviewer/adjudication"],
    ], widths=[2250, 7110], font_size=8.4)

    add_h(doc, "9.9. Entity resolution và quality gates", 2)
    add_p(doc, "customer_id nội bộ là canonical ID. Business registration/tax ID, LEI và vendor ID là external identifiers có source/version. Exact stable-ID match được ưu tiên; fuzzy name/address chỉ tạo candidate. High-impact merge hoặc switch phải được xác nhận. Tuyệt đối không join hai doanh nghiệp chỉ bằng tên viết tắt.")
    add_table(doc, ["Quality dimension", "Kiểm tra", "Khi fail"], [
        ["Completeness/validity", "Required fields, type, enum, unit, effective range", "Quarantine hoặc pending information"],
        ["Uniqueness/consistency", "Keys, duplicate, CRM–DMS–registry conflicts", "Conflict report/pending review"],
        ["Freshness", "TTL/update/effective date", "Mark stale; block time-sensitive decision"],
        ["Provenance", "Source record/location/hash/version", "Không publish claim/index/rule"],
        ["Access", "ACL/customer/employee scope", "Fail closed"],
        ["Extraction quality", "OCR/table/heading/citation samples", "Human review hoặc loại chunk"],
        ["Drift", "Schema, coverage, value/rank distribution", "Alert, canary, re-index/recalibrate"],
    ], widths=[2300, 4200, 2860], font_size=8.3)

    add_h(doc, "9.10. MVP data pack và Definition of Done", 2)
    add_bullets(doc, [
        "10 synthetic companies thuộc 4–5 segment/industry, stable IDs và permission scopes.",
        "8–12 products; 20–30 product/legal/SOP documents có version/effective dates, gồm superseded/conflict cases.",
        "5–10 blocking/warning/missing rules có source mapping.",
        "50 intent conversations có workspace context, abbreviations, corrections và multi-intent.",
        "40 RAG queries; 40 eligibility cases; 40 E2E; 25 security; 20 reliability scenarios.",
        "Một official/vendor adapter POC ở shadow mode; không bắt buộc cho offline MVP.",
    ])
    add_p(doc, "Data Definition of Done: 100% source phục vụ có valid Source Card, owner/purpose/lineage/version; mọi Product/Legal important claim trace được Gold → Silver → raw/official source; unauthorized/unlicensed source exposure = 0; stale policy used for time-sensitive decision = 0; high-risk entity merge false positive = 0; ingest report tái lập được.")

    add_h(doc, "9.11. Nguồn tham khảo cho market scan", 2)
    add_table(doc, ["Nguồn", "URL"], [
        ["Cổng đăng ký doanh nghiệp quốc gia", "https://dangkykinhdoanh.gov.vn/vn/pages/trangchu.aspx"],
        ["CSDL quốc gia về văn bản pháp luật", "https://vbpl.vn/Pages/portal.aspx"],
        ["CIC", "https://cic.gov.vn/"],
        ["National Statistics Office – NSDP", "https://nsdp.nso.gov.vn/"],
        ["FiinGroup", "https://fiingroup.vn/"],
        ["GLEIF data/API", "https://www.gleif.org/en/lei-data/access-and-use-lei-data"],
        ["OpenCorporates API", "https://api.opencorporates.com/"],
        ["UN sanctions", "https://main.un.org/securitycouncil/en/content/un-sc-consolidated-list"],
        ["OFAC Sanctions List Service", "https://ofac.treasury.gov/sanctions-list-service"],
        ["World Bank Debarred Firms", "https://www.worldbank.org/en/projects-operations/procurement/debarred-firms"],
        ["LSEG Risk Intelligence", "https://www.lseg.com/en/risk-intelligence"],
        ["Luật Bảo vệ dữ liệu cá nhân 91/2025/QH15", "https://vbpl.vn/TW/Pages/ivbpq-toanvan.aspx?ItemID=179252"],
    ], widths=[3600, 5760], font_size=7.6)
    add_p(doc, "Luật Bảo vệ dữ liệu cá nhân số 91/2025/QH15 có hiệu lực từ 01/01/2026; Legal/Privacy phải thẩm định cụ thể từng source và flow. Dữ liệu public không mặc nhiên được tự do ingest/reuse; PEP/adverse media và hồ sơ khách hàng cần purpose, access, review, retention và false-positive remediation rõ ràng.")

    add_h(doc, "9.12. Vertical slice dữ liệu cho case bán hàng doanh nghiệp", 2)
    add_p(doc, "MVP không bắt đầu từ catalog sản phẩm. Nó bắt đầu từ một câu chuyện bán hàng có đủ dữ liệu để mọi module cùng làm việc: RM đang mở hồ sơ Công ty ABC, biết doanh nghiệp có 500 nhân sự, nhiều nhà cung cấp và dòng tiền phân tán; khách hàng nói ngắn gọn rằng muốn trả lương, trả nhà cung cấp và cần vốn cho mùa cao điểm. Context Engine bổ sung dữ liệu đã có, Intent Engine chuẩn hóa nhu cầu, Product tạo bundle, Legal kiểm tra điều kiện, Operations chuẩn bị checklist và Evidence kiểm tra mọi claim trước khi RM duyệt.")
    add_table(doc, ["Thuộc tính case hero", "Giá trị synthetic MVP", "Agent/module sử dụng"], [
        ["Định danh", "case_id=CORP-DEMO-001; customer_id=COMP-ABC; rm_id=RM-999", "Tất cả module, audit, approval"],
        ["Doanh nghiệp", "Công ty CP ABC Việt Nam; sản xuất hàng tiêu dùng; 500 nhân sự; doanh thu 120 tỷ VND/năm", "Context, Intent, Product, Legal"],
        ["Vận hành tiền", "6 tài khoản tại 3 đơn vị; dòng tiền phân tán; khoảng 1.200 lệnh chi nhà cung cấp/tháng", "Intent, Product, Operations"],
        ["Hệ thống", "ERP có khả năng xuất file; API integration chưa xác nhận", "Product, Operations; API Banking là alternative"],
        ["Nhu cầu", "Payroll + supplier payment + cash management + working capital 20 tỷ/6 tháng", "Intent, Planner, Product"],
        ["Hồ sơ sẵn có", "Đăng ký doanh nghiệp hợp lệ", "Legal, Evidence"],
        ["Hồ sơ thiếu", "UBO và BCTC năm gần nhất", "Legal blocking cho nhánh tín dụng"],
        ["Kết quả kỳ vọng", "Giao dịch: ready_to_prepare; tín dụng: pending_information; email/checklist ở dạng draft", "Workflow, Operations, Approval"],
    ], widths=[2200, 4300, 2860], font_size=8.4)
    add_callout(doc, "Luật nhất quán dữ liệu", "Mọi record phát sinh từ case này phải dùng cùng canonical IDs. Product không được đổi customer segment; Legal không tự tạo hồ sơ chưa có; Operations không được coi nhánh tín dụng đã sẵn sàng; Evidence không được trích một policy khác version. Nếu một agent thiếu dữ liệu, nó ghi missing field vào shared state thay vì tự bịa.", fill=RED_LIGHT, border=RED)

    add_h(doc, "9.13. Form F-01 — RM Workspace và context nhân viên", 2)
    add_p(doc, "Form này được hệ thống tự lắp từ SSO/IAM, màn hình đang mở, CRM và lịch sử case. RM chỉ sửa khi phát hiện sai; không phải nhập lại mỗi lần. Mỗi field lưu value, source, observed_at, confidence và confirmed.")
    add_table(doc, ["Nhóm", "Field bắt buộc", "Nguồn ưu tiên", "MVP hero value"], [
        ["Nhân viên", "employee_id, role, branch_id, permission_scopes", "SSO/IAM mock", "RM-999; Corporate RM; BR-HN-01; case:read/draft/approve"],
        ["Workspace", "active_customer_id, active_case_id, screen, selected_product_ids", "UI session", "COMP-ABC; CORP-DEMO-001; sales-case; []"],
        ["Công việc", "current_task, due_at, previous_artifact_ids", "CRM/task mock", "discovery; ngày demo; chưa có artifact"],
        ["Hội thoại", "last_user_message, last_confirmed_intent, answered_slots", "Conversation state", "Yêu cầu bán hàng; chưa confirm; slots từ CRM"],
        ["Freshness", "source_updated_at, loaded_at, stale_after, is_stale", "Context service", "realtime/session cho workspace; TTL theo source"],
        ["Quyền", "customer_scope_allowed, document_scope_allowed", "IAM policy", "true cho COMP-ABC; fail-closed nếu không xác minh"],
    ], widths=[1500, 3100, 2200, 2560], font_size=7.9)

    add_h(doc, "9.14. Form F-02 — Sales discovery cho khách hàng doanh nghiệp", 2)
    add_p(doc, "Đây là form nghiệp vụ quan trọng nhất cho Intent và Product matching. Form phải hỗ trợ ba trạng thái: known (đã có), inferred (suy luận có provenance) và unknown (chưa biết). Unknown không đồng nghĩa phải hỏi ngay; chỉ hỏi nếu field đó chặn quyết định hiện tại.")
    add_table(doc, ["Khối dữ liệu", "Trường cần thu thập", "Bắt buộc ở MVP", "Cách điền hiệu quả"], [
        ["Định danh doanh nghiệp", "tax_code, legal_name, industry, legal_type, operating_years", "tax_code, name, industry", "CRM trước; registry dùng verify/enrichment"],
        ["Quy mô", "employees_count, annual_revenue, branches, accounts_count", "employees_count, revenue", "CRM/BCTC; không hỏi lại khi còn fresh"],
        ["Chi lương", "payroll_headcount, payroll_value, pay_day, employee_account_ratio", "headcount", "Dùng profile làm default; hỏi pay_day sau khi chọn sản phẩm"],
        ["Nhà cung cấp", "supplier_count, payments_per_month, file_or_api, approval_levels", "payments_per_month", "Cho phép nêu trong câu tự nhiên; extractor chuẩn hóa"],
        ["Dòng tiền", "accounts/branches, concentration, sweeping need, reporting need", "cash_flow_status", "CRM + discovery; tạo Cash Management signal"],
        ["Nhu cầu vốn", "amount, tenor, purpose, expected_draw_date, collateral, contracts/invoices", "amount, tenor, purpose", "Thiếu BCTC/UBO vẫn cho tạo sơ bộ nhưng block approval"],
        ["Hệ thống", "ERP/HRM, API readiness, file formats, technical contact", "không bắt buộc", "Chỉ hỏi sau khi API/File product trở thành candidate"],
        ["Ưu tiên", "urgency, target_date, decision_makers, contact channel", "urgency", "RM xác nhận trước khi soạn task/email"],
    ], widths=[1800, 3300, 1750, 2510], font_size=7.7)
    add_code(doc, '{\n  "case_id": "CORP-DEMO-001",\n  "company": {"customer_id": "COMP-ABC", "employees_count": 500, "annual_revenue_vnd": 120000000000},\n  "sales_signals": {"supplier_payments_per_month": 1200, "cash_flow": "distributed", "erp": "file_export"},\n  "funding_need": {"amount_vnd": 20000000000, "tenor_months": 6, "purpose": "seasonal_working_capital"},\n  "unknown_fields": ["pay_day", "approval_matrix", "api_readiness"]\n}')

    add_h(doc, "9.15. Form F-03 đến F-08 — Contract dữ liệu cho từng agent", 2)
    add_p(doc, "Các form dưới đây không phải prompt riêng lẻ. Đây là hợp đồng dữ liệu giữa các module. Mỗi output có schema_version, case_id, generated_at, provenance và validation_status; chỉ shared state là nơi hợp nhất trạng thái cuối.")
    add_table(doc, ["Form / owner", "Input liên quan trực tiếp đến case", "Output bắt buộc", "Không được phép"], [
        ["F-03 Intent / Intent Engine", "RM context + discovery + message + previous confirmations", "primary_intent; secondary_intents; slots; missing_decision_fields; confidence; provenance", "Tự thêm product hoặc kết luận eligibility"],
        ["F-04 Product / Product Agent", "IntentResult + company segment/size + controlled catalog/policy", "recommended_bundle; alternatives; score components; prerequisites; evidence_ids", "Đề xuất product ngoài catalog; ghi eligible=true"],
        ["F-05 Eligibility / Legal", "Company legal profile; UBO/KYC; documents; selected products; versioned rules", "status theo từng product/branch; failed_checks; missing_documents; rule/evidence IDs", "Downgrade blocking rule; tự coi unknown là pass"],
        ["F-06 Operations / Operations", "Approved candidate bundle + legal branch status + SOP/SLA + existing tasks", "checklist; task drafts; CRM case draft; customer email draft; dedup keys", "Tạo task/email thật; tạo trùng artifact"],
        ["F-07 Evidence / Validator", "Claims của Product/Legal/Ops + retrieved chunks + source metadata", "claim_supported; source/version/section; conflict/stale flags; validation report", "Xác nhận claim không có source hoặc source hết hiệu lực"],
        ["F-08 Approval / Executor", "Validated payload; actor/scope; risk; payload hash; idempotency key", "preview; approval event; execution receipt; audit", "Thực thi khi token sai payload, hết hạn hoặc đã dùng"],
    ], widths=[1800, 3100, 2900, 1560], font_size=7.5)

    add_h(doc, "9.16. Mẫu dữ liệu chi tiết cho Product, Legal và Operations", 2)
    add_h(doc, "Product Catalog record", 3)
    add_code(doc, '{\n  "product_id": "SYNTH-PROD-BULK-PAYMENT", "product_name": "Chi hộ nhà cung cấp",\n  "category": "payments", "target_segments": ["SME", "large_corporate"],\n  "supported_needs": ["supplier_payment"], "features": ["file_payment", "multi_level_approval"],\n  "prerequisites": ["corporate_payment_account"], "required_documents": ["service_agreement", "business_registration"],\n  "compatible_products": ["SYNTH-PROD-CASH-MGMT"], "exclusion_conditions": [],\n  "source_document_ids": ["SYNTH-DOC-PRODUCT-004"], "document_version": "1.0",\n  "effective_date": "2026-01-01", "status": "active", "data_label": "SYNTHETIC DEMO DATA"\n}')
    add_h(doc, "Eligibility Rule record", 3)
    add_code(doc, '{\n  "rule_id": "SYNTH-RULE-WC-FS-001", "product_id": "SYNTH-PROD-WORKING-CAPITAL",\n  "field": "financial_reports.has_recent", "operator": "eq", "expected": true,\n  "severity": "blocking", "on_unknown": "pending_information",\n  "required_evidence_types": ["financial_statement"],\n  "source_document_id": "SYNTH-DOC-CREDIT-001", "source_section": "5.2",\n  "version": "1.0", "effective_from": "2026-01-01", "data_label": "SYNTHETIC DEMO DATA"\n}')
    add_h(doc, "Operations SOP record", 3)
    add_code(doc, '{\n  "workflow_id": "SYNTH-SOP-CORP-SALES-001", "product_id": "SYNTH-PROD-BULK-PAYMENT",\n  "step_id": "OPS-03", "sequence": 3, "precondition": "eligibility_status != failed",\n  "task_template": "Xác nhận ma trận phê duyệt và định dạng file thanh toán",\n  "owner_role": "RM", "sla_hours": 8, "approval_required": false,\n  "dedup_key_template": "{case_id}:{product_id}:{step_id}",\n  "source_document_id": "SYNTH-DOC-SOP-002", "version": "1.0"\n}')

    add_h(doc, "9.17. Sinh dữ liệu synthetic hiệu quả bằng scenario graph", 2)
    add_p(doc, "Cách sinh dữ liệu hiệu quả nhất cho MVP là scenario-first, constraint-driven và reproducible. Không dùng LLM để tạo tự do hàng nghìn record ngay từ đầu. Trước tiên định nghĩa archetype doanh nghiệp, nhu cầu, policy/rules và expected outcome; sau đó generator tạo các bảng liên quan bằng seed cố định. LLM chỉ hỗ trợ paraphrase câu nói của RM và tạo biến thể tài liệu; code deterministic giữ ID, số tiền, trạng thái và ground truth.")
    add_bullets(doc, [
        "Bước 1 — Chọn archetype: manufacturing-500-staff, retail-multi-outlet, logistics-import-export.",
        "Bước 2 — Gắn need graph: payroll, supplier_payment, cash_management, working_capital.",
        "Bước 3 — Chọn product bundle và tạo prerequisite/rule graph tương ứng.",
        "Bước 4 — Áp dụng missingness có chủ đích: thiếu UBO/BCTC, stale policy, conflicting representative hoặc tool timeout.",
        "Bước 5 — Sinh conversation variants: đủ dấu/không dấu, viết tắt, câu ngắn, correction, multi-intent và out-of-scope.",
        "Bước 6 — Tạo ground truth trước output model: expected intents, product IDs, blocking rules, missing docs, workflow status và allowed actions.",
        "Bước 7 — Validate referential integrity, schema, range, effective dates, ACL và evidence linkage.",
        "Bước 8 — Đóng gói manifest gồm seed, generator_version, source template hashes và expected metrics.",
    ], numbered=True)
    add_table(doc, ["Tập dữ liệu MVP ngày mai", "Số lượng tối thiểu", "Mục tiêu"], [
        ["Product catalog", "6 sản phẩm", "Account, eBanking, Payroll, Supplier Payment, Cash Management, Working Capital"],
        ["Policy/SOP documents", "8–10 tài liệu ngắn", "Mỗi claim quan trọng có source/version/section"],
        ["Company profiles", "3 hồ sơ", "Một hero case + hai regression cases"],
        ["Sales conversations", "20 biến thể", "Keyword, không dấu, correction, multi-intent"],
        ["Eligibility scenarios", "12 cases", "Pass, block, missing, stale, conflict"],
        ["E2E golden cases", "10 cases", "6 normal, 2 edge, 1 adversarial, 1 tool failure"],
        ["Security cases", "5 cases", "Injection, wrong RM, approval tamper, replay, PII log"],
    ], widths=[3000, 1800, 4560], font_size=8.4)
    add_callout(doc, "Phân bổ dữ liệu khuyến nghị", "60% happy/near-happy path để demo ổn định; 25% missing/conflict để chứng minh hệ thống biết dừng; 15% adversarial/failure để chứng minh guardrail. Mỗi high-risk negative case phải có expected status và expected forbidden action, không chỉ expected text.", fill=GREEN_LIGHT, border=GREEN)

    add_h(doc, "9.18. Form F-09 — Synthetic Scenario Specification", 2)
    add_table(doc, ["Field", "Ý nghĩa", "Ví dụ cho case hero"], [
        ["scenario_id / seed", "ID và seed tái lập", "SCN-CORP-SALES-001 / 20260718"],
        ["archetype", "Mẫu doanh nghiệp", "manufacturing_500_staff"],
        ["input_message", "Câu RM nhập", "Khách cần chi lương, trả NCC và 20 tỷ vốn mùa cao điểm"],
        ["context_overrides", "Dữ liệu đã có trong workspace/CRM", "customer_id, 500 staff, 120 tỷ revenue, distributed cash flow"],
        ["missingness", "Field/document cố ý thiếu", "UBO, recent BCTC"],
        ["expected_intents", "Ground truth intent", "payroll; supplier_payment; cash_management; working_capital"],
        ["expected_bundle", "Ground truth product IDs", "6 controlled synthetic products"],
        ["expected_branch_status", "Kết quả theo nhánh", "transaction=ready_to_prepare; credit=pending_information"],
        ["expected_artifacts", "Checklist/draft/task", "3 missing-info tasks + 1 email draft"],
        ["forbidden_actions", "Hành động không được xảy ra", "approve_credit; send_email; create_live_crm_without_approval"],
        ["evidence_expectation", "Nguồn bắt buộc", "product/rule/SOP IDs và version khớp"],
    ], widths=[2450, 3000, 3910], font_size=8.1)

    add_h(doc, "9.19. Đánh giá dữ liệu thị trường theo tính phù hợp và khả dụng", 2)
    add_p(doc, "Market scan chỉ trả lời nguồn nào tồn tại và có thể thử nghiệm. Nó không tự cấp quyền sử dụng. Trước khi production ingest, từng nguồn phải qua Source Card, legal basis/license, owner, access method, data minimization, retention và quality POC. Trong MVP ngày mai, nguồn thị trường chỉ được mô phỏng hoặc dùng làm thiết kế adapter; không phụ thuộc kết nối mạng khi demo.")
    add_table(doc, ["Loại dữ liệu hiện có", "Ví dụ nguồn thị trường", "Khả dụng thực tế", "Vai trò đúng trong solution"], [
        ["Đăng ký pháp nhân Việt Nam", "Cổng thông tin quốc gia về đăng ký doanh nghiệp", "Public search và information services; bulk/API cần xác minh quyền/kênh", "Verify tên, mã số, địa chỉ, đại diện, tình trạng; không thay hồ sơ KYC nội bộ"],
        ["Văn bản pháp luật", "CSDL quốc gia về văn bản pháp luật; SBV publications", "Public web/documents; cần theo dõi hiệu lực và văn bản thay thế", "Legal RAG/reference; rule quan trọng vẫn được Legal owner mã hóa và ký nhận"],
        ["Thống kê ngành/vĩ mô", "NSO/NSDP; SBV statistics", "Public aggregate data", "Context và benchmark; không quyết định eligibility của một doanh nghiệp"],
        ["Legal entity toàn cầu", "GLEIF API/Golden Copy; OpenCorporates", "GLEIF mở và cập nhật thường xuyên; OpenCorporates theo license/plan", "Entity resolution, ownership enrichment, cross-border context"],
        ["Danh sách chế tài", "UN Consolidated List; OFAC SLS", "Public XML/HTML/PDF; cập nhật biến động", "Screening candidate; cần fuzzy matching, threshold và human review false positive"],
        ["Thông tin tín dụng", "CIC", "Không phải open data; chỉ qua quyền/kênh nghiệp vụ được cấp", "Live tool cho credit assessment; không lưu vào RAG chung"],
        ["Dữ liệu doanh nghiệp thương mại", "FiinGroup, D&B, Moody's/Orbis", "Licensed; coverage/field/freshness/cost phụ thuộc hợp đồng", "Financial/company enrichment ở shadow mode trước khi dùng quyết định"],
        ["KYC/AML/adverse media", "LSEG Risk Intelligence và vendor tương đương", "Licensed; cần POC recall/precision, explainability và review queue", "Risk signal; không auto-reject chỉ từ fuzzy match"],
        ["Dữ liệu nội bộ ngân hàng", "CRM, Core, DMS, transaction, task, product/policy/SOP", "Giá trị cao nhất nhưng phụ thuộc owner, IAM, API và chất lượng", "Nguồn quyết định chính cho context, product, eligibility và operations"],
    ], widths=[1900, 2300, 2250, 2910], font_size=7.3)
    add_p(doc, "Nguồn chính thức được kiểm tra tại thời điểm 17/07/2026 cho thấy: Cổng đăng ký doanh nghiệp hỗ trợ tra cứu và dịch vụ thông tin; GLEIF cung cấp API và Golden Copy/delta data, trong đó Level 1 trả lời “who is who” và Level 2 hỗ trợ quan hệ sở hữu; UN và OFAC cung cấp sanctions data ở định dạng máy đọc. Những nguồn này phù hợp để verify/enrich/screen, nhưng quyết định bán sản phẩm và eligibility vẫn cần dữ liệu nội bộ, policy hiện hành và human governance.")

    add_h(doc, "9.20. Pipeline chuẩn bị và xử lý dữ liệu theo từng agent", 2)
    add_code(doc, "Source registration + owner/license/purpose\n→ Raw/Quarantine (immutable hash, malware/OCR/encoding checks)\n→ Silver (canonical IDs, normalized units, dedup, entity resolution, ACL)\n→ Gold by agent (Product catalog / Rule registry / SOP / Context views / Eval labels)\n→ Serving (API, rule engine, sparse+dense index, feature/context store)\n→ Trace + feedback + correction + version retirement")
    add_table(doc, ["Agent/module", "Gold dataset", "Xử lý đặc thù", "Quality gate trước serving"], [
        ["Context/Intent", "employee_workspace_snapshot; customer_360; conversation_state", "Freshness, precedence, conflict, minimization, PII masking", "No cross-case leak; required IDs; source/confidence/confirmed"],
        ["Product", "product_master; product_policy_chunks; compatibility graph", "Structure-aware parse, version/effective dates, controlled vocabulary, hybrid index", "Active version only; product/source linkage 100%; retrieval golden pass"],
        ["Legal/Eligibility", "rule_registry; legal_chunks; KYC/UBO/document status", "Deterministic operators, severity, on_unknown, validity windows, watchlist matching", "Unsafe pass=0; every blocking rule has evidence and owner"],
        ["Operations", "workflow/SOP; SLA/calendar; templates; existing task fingerprints", "Step order, precondition, owner, dedup key, idempotency", "No duplicate task; template variables complete; no send permission"],
        ["Evidence", "evidence_items; source manifests; claim-source links", "Quote/section validation, staleness/conflict detection", "Important claim support=100%; stale/conflict blocks"],
        ["Approval/Audit", "payload snapshots; approval events; execution receipts", "Hash, nonce, expiry, actor/scope, sanitized logs", "Payload equality; one-time use; replay blocked"],
        ["Evaluation", "golden scenarios; expected outputs/actions; adjudication", "Seed/version, difficulty, risk, provenance, reviewer agreement", "Schema pass; label review; no train/test leakage"],
    ], widths=[1700, 2500, 3000, 2160], font_size=7.6)

    add_h(doc, "9.21. Data acceptance checklist cho MVP", 2)
    add_bullets(doc, [
        "Case hero chạy offline từ create → pending_information → bổ sung UBO/BCTC → pending_approval → mock execute.",
        "Mọi agent output chứa đúng case_id/customer_id và validate được bằng schema; không orphan record.",
        "Sáu sản phẩm synthetic có ID, supported needs, prerequisites, source, version, effective date và data label.",
        "Nhánh transaction không bị mất chỉ vì nhánh working capital bị block.",
        "Mỗi blocking check chỉ ra field/document thiếu và source rule; unknown không được biến thành pass.",
        "Email/task chỉ là draft trước approval; approval token không cho phép payload đã bị sửa.",
        "Golden scenarios tái lập bằng seed; generator và dataset manifest được version hóa.",
        "Tất cả nguồn thị trường có availability/decision role rõ; source chưa có license không được ingest production.",
    ])

    add_h(doc, "10. PRODUCT KNOWLEDGE, INGESTION VÀ HYBRID RAG", 1)
    add_h(doc, "10.1. Giới hạn cần khắc phục", 2)
    add_p(doc, "Baseline hiện tại dùng hash embedding deterministic và catalog in-memory. Cách này phù hợp demo nhưng chưa đủ cho PDF/Word/Excel thật, versioning, ACL, hiệu lực chính sách, persistent vector index và đánh giá retrieval. V2 phải có ingestion report và index manifest; không được gọi một danh sách hard-code là production RAG.")

    add_h(doc, "10.2. Nguồn và ingestion", 2)
    add_table(doc, ["Nguồn", "Định dạng", "Metadata bắt buộc"], [
        ["Product catalog", "PDF/Excel/DB", "product_id, family, segment, version, owner"],
        ["Product policy", "PDF/Word", "effective_from/to, status, rule links, ACL"],
        ["Fee/limit tables", "Excel/table", "currency, unit, version, effective date"],
        ["FAQ/Sales guide", "Docs", "audience, scope, product mapping"],
    ], widths=[2300, 1800, 5260], font_size=8.8)
    add_code(doc, "File/API → SHA-256 + document/version → parser/OCR router\n→ text/table extraction → Unicode cleaning + quality checks\n→ structure-aware chunking → metadata/ACL enrichment\n→ dense embedding + sparse index → manifest + ingest report")
    add_p(doc, "Chunk phải giữ section path, page, product ID, effective date, active flag, access scope, content hash và parent/neighbor references. Bảng được lưu thành summary + row chunks nhưng luôn giữ header và unit; eligibility rule không tách khỏi product ID.")

    add_h(doc, "10.3. Retrieval và matching", 2)
    add_code(doc, "query normalization + resolved slots\n→ ACL/effective-date/segment filters\n→ dense top-20 + sparse top-20\n→ weighted fusion (khởi tạo 0.6/0.4, tune bằng eval)\n→ rerank + dedup + source diversity\n→ threshold/OOS gate → top 3–5 chunks")
    add_p(doc, "RAG chỉ tạo candidates. Product Matcher tính intent fit, segment fit, size/revenue fit và workflow signal; prerequisites thiếu được trừ điểm hoặc hiển thị riêng. Trường eligible phải để unknown cho tới Eligibility Engine. Không recommendation nào được phép sử dụng product name ngoài controlled catalog.")

    add_h(doc, "11. ELIGIBILITY, LEGAL VÀ COMPLIANCE", 1)
    add_h(doc, "11.1. Nguyên tắc phân quyền quyết định", 2)
    add_p(doc, "Rule deterministic sở hữu kết quả đạt/không đạt/pending cho điều kiện đã mã hóa. Legal RAG cung cấp điều khoản, phiên bản và giải thích. Live tools đọc KYC/UBO/watchlist khi được phép. LLM không được downgrade severity, không tự kết luận “đủ điều kiện” và không được bỏ qua input stale.")
    add_code(doc, "Permission/sanction hard block\n→ legal/regulatory blocking\n→ product eligibility blocking\n→ missing required information\n→ warning/advisory\n→ LLM explanation grounded by evidence")

    add_h(doc, "11.2. Semantics kết quả", 2)
    add_table(doc, ["Status", "Điều kiện"], [
        ["passed", "Tất cả blocking rules pass và dữ liệu bắt buộc còn fresh"],
        ["failed", "Có rule loại trừ rõ ràng"],
        ["pending_information", "Thiếu hoặc stale input/document bắt buộc"],
        ["pending_review", "Policy conflict, PEP/AML, dữ liệu mâu thuẫn hoặc ngoại lệ pháp lý"],
    ], widths=[2400, 6960], font_size=9)
    add_p(doc, "Trong case ABC, thiếu UBO và BCTC chỉ chặn nhánh Working Capital. Payroll/Cash Management không bị loại bỏ nếu điều kiện riêng của chúng thỏa. Product vẫn hiển thị “blocked/pending” cùng lý do thay vì bị xóa khỏi phương án, giúp RM hiểu trade-off và thông tin cần bổ sung.")

    add_h(doc, "11.3. Failure policy", 2)
    add_bullets(doc, [
        "Rule registry không tải được: fail closed cho eligibility.",
        "Legal index lỗi: chỉ tiếp tục nếu rule/source cached còn hiệu lực; nếu không chuyển pending_review.",
        "KYC timeout: không được trả passed; chuyển pending_review/pending_information.",
        "Hai policy active mâu thuẫn: hiển thị cả hai nguồn, dừng để người có thẩm quyền review.",
        "Malformed blocking rule: quarantine + alert, không được silently ignore.",
    ])

    add_h(doc, "12. WORKFLOW ORCHESTRATION, RETRY VÀ PARTIAL RESUME", 1)
    add_h(doc, "12.1. Node contract", 2)
    add_table(doc, ["Node", "Đọc", "Ghi / outcome"], [
        ["collect_context", "request/session", "ContextSnapshot hoặc access/availability error"],
        ["extract_intent / resolve_slots", "message + minimized context", "IntentResult, confidence, clarification"],
        ["route_complexity / plan_tasks", "intent", "route, typed DAG, dependency"],
        ["retrieve_products", "intent/context", "product candidates + evidence"],
        ["evaluate_eligibility", "product/context/docs", "rule results + missing/blocking"],
        ["validate_evidence", "claims/evidence", "valid/invalid flags; re-retrieve/review"],
        ["prepare_operations / dedup", "validated results + existing artifacts", "drafts + reuse/update/create decision"],
        ["await_approval / execute", "frozen payload + token", "approved action result + audit"],
    ], widths=[2500, 2650, 4210], font_size=8.4)

    add_h(doc, "12.2. Routing và giới hạn vòng lặp", 2)
    add_p(doc, "Yêu cầu chỉ đọc, một intent, không eligibility rủi ro cao và context đủ có thể đi fast path. Multi-intent, credit/KYC, missing-information loop hoặc draft/write phải đi complex route. DAG phải kiểm unknown dependency và cycle; max adaptive loops = 3. Planner không có quyền gọi business write tool.")

    add_h(doc, "12.3. Retry và idempotency", 2)
    add_table(doc, ["Failure", "Retry", "Giới hạn / điều kiện"], [
        ["Model timeout/5xx", "Có", "2 lần, exponential backoff"],
        ["Schema parse", "Repair", "1 lần, sau đó fallback/typed error"],
        ["Read tool timeout", "Có nếu an toàn", "1 lần, cache fallback"],
        ["Write tool timeout", "Chỉ có idempotency", "Query status trước retry"],
        ["Permission denied", "Không", "Fail closed"],
        ["Invalid evidence", "Không retry action", "Re-retrieve một lần hoặc review"],
    ], widths=[2600, 2200, 4560], font_size=8.8)

    add_h(doc, "12.4. Impact graph", 2)
    add_table(doc, ["Thay đổi", "Resume", "Giữ lại"], [
        ["UBO/BCTC mới", "Eligibility → Evidence → Operations", "Context, intent, product"],
        ["Đổi customer", "Context → toàn bộ downstream", "Audit"],
        ["Đổi mục tiêu", "Intent → toàn bộ downstream", "Employee context"],
        ["Catalog version mới", "Product → Eligibility → downstream", "Context/intent"],
        ["RM sửa email", "Approval payload", "Analysis results"],
    ], widths=[2200, 3700, 3460], font_size=8.8)

    add_h(doc, "13. OPERATIONS: CHECKLIST, ARTIFACT REUSE VÀ DEDUP", 1)
    add_h(doc, "13.1. Output chỉ là draft", 2)
    add_p(doc, "Operations nhận intent/context đã validate, product recommendation có evidence, eligibility result, existing case/task/artifacts và SOP version. Module tạo decision brief, checklist, customer message draft, CRM case draft và task drafts; không được tạo side effect.")
    add_h(doc, "13.2. Checklist engine", 2)
    add_p(doc, "Checklist là union có giải thích của product prerequisites, legal missing documents, KYC/UBO, SOP và context-specific requirements. Dedup dùng controlled document taxonomy, không chỉ string equality. Mỗi item lưu document_type_id, status, reason, product/rule/evidence IDs và existing document reference.")

    add_h(doc, "13.3. Drafting an toàn", 2)
    add_bullets(doc, [
        "Dùng template trước; LLM chỉ cải thiện văn phong từ structured verified fields.",
        "Không thêm phí, lãi suất, hạn mức, deadline hoặc điều kiện không có source.",
        "Recipient từ CRM chỉ là candidate và phải được RM verify trước send.",
        "RM edit tạo version/content hash mới và làm approval cũ mất hiệu lực.",
        "Email không cam kết phê duyệt tín dụng; nêu rõ hồ sơ cần bổ sung và mục đích liên hệ.",
    ])

    add_h(doc, "13.4. Dedup key và reuse policy", 2)
    add_code(doc, "org + customer_id + case/business_request_id + task_type\n+ product_id? + workflow_step + normalized_subject_hash")
    add_table(doc, ["Artifact hiện có", "Input/version", "Hành động"], [
        ["Active task", "Giống", "Reuse/attach"],
        ["Active task", "Khác", "Update nếu cho phép; nếu không tạo linked revision"],
        ["Completed task", "Còn validity", "Reuse result"],
        ["Completed task", "Stale", "Tạo mới với supersedes link"],
        ["Email draft chưa gửi", "Cùng purpose", "Update draft hiện có"],
        ["CRM case active", "Cùng request", "Append/update, không create duplicate"],
    ], widths=[2800, 2300, 4260], font_size=8.8)

    add_h(doc, "14. EVIDENCE, GUARDRAILS, APPROVAL VÀ EXECUTION", 1)
    add_h(doc, "14.1. Defense in depth", 2)
    add_bullets(doc, [
        "Authentication/session và RBAC/ABAC trước retrieval.",
        "File type/size/malware, prompt injection và PII minimization ở input.",
        "ACL/effective date filter trước khi chunk đến model.",
        "Schema validation và deterministic exact-match cho số, phí, limit, unit.",
        "Tool allowlist theo caller/module, risk, scope, approval và idempotency.",
        "Approval token gắn case, approver, permissions, payload hash, expiry, nonce và one-time use.",
        "Executor load latest state và verify lại evidence/blocking/permission trước side effect.",
    ])

    add_h(doc, "14.2. Evidence validation", 2)
    add_p(doc, "Validator kiểm source identity/version/effective status, quote presence, deterministic value/unit match và semantic support. Claim không được hỗ trợ bị loại khỏi output khách hàng, đánh hallucination_flag, re-retrieve một lần và sau đó chuyển review/failure. Numeric claims không được pass chỉ bằng semantic similarity.")

    add_h(doc, "14.3. Approval integrity", 2)
    add_code(doc, "verify auth/session → load latest state\n→ verify signature/expiry/nonce/one-time use\n→ recompute payload hash\n→ verify evidence + no blocking + permissions\n→ acquire idempotency lock → call adapter\n→ reconcile uncertain outcome → persist audit → consume token")
    add_callout(doc, "Nguyên tắc bất biến", "AI chỉ phân tích, đề xuất và soạn nháp. Tạo case/task hoặc gửi phản hồi ra ngoài chỉ được thực hiện khi RM (hoặc cấp có thẩm quyền theo matrix) phê duyệt đúng payload. Chỉnh một ký tự trong payload sau approval cũng phải làm token cũ mất hiệu lực.", fill=RED_LIGHT, border=RED)

    add_h(doc, "15. API V2 VÀ RM WORKSPACE", 1)
    add_h(doc, "15.1. API contract", 2)
    add_table(doc, ["Method", "Endpoint", "Mục đích"], [
        ["GET", "/api/v2/context/current", "Context employee/workspace hiện tại"],
        ["POST", "/api/v2/context/resolve", "Assemble context cho message"],
        ["POST/GET", "/api/v2/cases", "Tạo analysis case / đọc state"],
        ["POST", "/api/v2/cases/{id}/messages", "Yêu cầu hoặc clarification answer"],
        ["POST", "/api/v2/cases/{id}/documents", "Đăng ký/upload metadata tài liệu"],
        ["PATCH", "/api/v2/cases/{id}/context", "Sửa context + reason + expected version"],
        ["POST", "/api/v2/cases/{id}/resume", "Server tính impacted nodes và resume"],
        ["POST", "/api/v2/cases/{id}/approval-preview", "Freeze payload + diff"],
        ["POST", "/api/v2/cases/{id}/approve|execute|reject", "HITL và action có kiểm soát"],
        ["GET", "/api/v2/cases/{id}/trace", "User-safe timeline; không hiển thị hidden CoT"],
    ], widths=[1000, 3500, 4860], font_size=8.4)
    add_p(doc, "Auth principal lấy từ session/token, không tin employee_id trong body. API dùng stable error codes, trace ID, ETag/state version và Idempotency-Key cho write. Client không được tự chọn bỏ qua evidence node khi resume.")

    add_h(doc, "15.2. Các panel trong UI", 2)
    add_bullets(doc, [
        "Context Header: RM/role, customer, active case, current step, product, missing info; có source/freshness tooltip.",
        "Intent Preview: “Hệ thống hiểu rằng…”, primary/sub-intents, resolved fields, assumptions và một clarification nếu blocking.",
        "Product/Evidence: candidates, score components, eligibility, source quote/version, blocked/pending reason.",
        "Operations: checklist, existing/reuse/update/create badge, email editor/version diff và SLA source.",
        "Approval: exact actions, target/recipient, payload diff, risk/evidence và approve/reject.",
        "Timeline: context loaded, intent resolved, retrieval/rules, draft reuse/update, approval/execution; không phơi chain-of-thought.",
    ])

    add_h(doc, "16. STORAGE, OBSERVABILITY VÀ RELIABILITY", 1)
    add_h(doc, "16.1. Persistent storage", 2)
    add_table(doc, ["Store/Table", "Mục đích"], [
        ["cases / case_state_versions", "Current index + immutable versioned state/hash"],
        ["workflow_tasks", "DAG node, dedup key, status, input/output hash"],
        ["context_values", "Field-level provenance, confidence, freshness"],
        ["artifacts", "Checklist/email/brief/case/task drafts theo version/hash"],
        ["approval_tokens / idempotency_records", "One-time approval và side-effect dedup"],
        ["audit_events", "Append-only hash chain, actor/action/state versions"],
        ["Vector DB", "Product/legal chunks + ACL/effective metadata"],
        ["DMS/Object store", "Source documents; không nhúng raw blob trong case state"],
        ["Redis/cache", "Context/retrieval/node caches theo TTL/version/scope"],
    ], widths=[3400, 5960], font_size=8.8)

    add_h(doc, "16.2. Observability tối thiểu", 2)
    add_bullets(doc, [
        "Trace ID xuyên API → context → intent → workflow → retrieval/rules/tool.",
        "Log JSON có event code, prompt/workflow/rule/index version và sanitized IDs; không log raw PII, token hoặc email nhạy cảm.",
        "Metrics cho context stale/conflict/auto-fill, intent schema/clarification, RAG hit/empty/latency, eligibility pending/block, resume/dedup, approval/action và cost.",
        "Timeout mọi network/model call; backoff chỉ cho safe reads, circuit breaker theo dependency, DLQ cho async job và reconciliation cho write timeout.",
        "Cache key luôn gồm version và permission scope; không reuse cross-customer/cross-scope.",
    ])

    add_h(doc, "16.3. SLO đề xuất", 2)
    add_table(doc, ["SLO", "Mục tiêu ban đầu"], [
        ["API read availability pilot", "99.5%"],
        ["P95 context assembly", "< 2 giây, không tính upstream unavailable"],
        ["P95 complete analysis", "< 30 giây"],
        ["Duplicate external write", "0"],
        ["High-risk alert emission", "< 1 phút"],
    ], widths=[4300, 5060], font_size=9)

    add_h(doc, "17. EVALUATION VÀ QUALITY GATES", 1)
    add_p(doc, "Không dùng một demo đẹp làm bằng chứng chất lượng. Hệ thống phải đo riêng context, intent, retrieval, eligibility, workflow, safety và end-to-end; deterministic metrics là gate chính, LLM-as-judge chỉ dùng cho clarity/tone hoặc semantic support khó xác định.")
    add_table(doc, ["Suite", "Quy mô tối thiểu pilot", "Mục đích"], [
        ["Intent conversations", "100", "Intent/entity/context/confidence/no-repeat"],
        ["Product RAG", "40", "Retrieval/citation/OOS/version/ACL"],
        ["Eligibility", "40", "Rules, missing, blocking, conflict"],
        ["E2E business", "40", "Complete journeys, artifact/action outcomes"],
        ["Adversarial/security", "25", "Injection, RBAC, tool, token, payload"],
        ["Reliability", "20", "Timeout, retry, cache, replay, concurrency"],
    ], widths=[2900, 2000, 4460], font_size=8.8)
    add_table(doc, ["Metric", "MVP gate", "Pilot gate"], [
        ["Contract-valid outputs", "100%", "100%"],
        ["Primary intent accuracy", "≥ 90%", "≥ 95%"],
        ["Multi-intent recall", "≥ 90%", "≥ 95%"],
        ["System slot auto-fill", "≥ 98%", "≥ 99%"],
        ["Unnecessary clarification", "< 10%", "< 5%"],
        ["Product Hit@5", "≥ 90%", "≥ 95%"],
        ["Citation correctness", "100% important claims", "100%"],
        ["Eligibility unsafe pass", "0%", "0%"],
        ["Missing-document recall", "≥ 95%", "100% high-risk target"],
        ["Duplicate task/action", "0%", "0%"],
        ["Correct resume selection", "≥ 90%", "≥ 95%"],
        ["Cross-scope leak", "0", "0"],
    ], widths=[3600, 2880, 2880], font_size=8.6)
    copy_source_table(doc, source, 28, "Bảng 10. Bộ metric nghiệp vụ/kỹ thuật của proposal ban đầu được giữ lại làm lớp bổ sung.", widths=[2400, 3500, 3460], font_size=8.3)

    add_h(doc, "18. KẾ HOẠCH BUILD THỐNG NHẤT CHO AI CODING", 1)
    add_h(doc, "18.1. Nguyên tắc contract-first", 2)
    add_p(doc, "JSON schemas là source of truth cho shared state, context, intent và tool contract. Nếu code cần khác contract, thay đổi phải đi cùng migration, tests và progress log. AI coding phải đọc INDEX → PROGRESS → build protocol → contract liên quan → module plan → acceptance trước khi báo hoàn thành.")
    add_bullets(doc, [
        "Không thêm field hoặc status rải rác trong code.",
        "Không bắt đầu task khi dependency chưa Done, trừ adapter mock có interface rõ.",
        "Mỗi task có unit, integration, metrics/log và security considerations.",
        "Không gọi một module là Done chỉ vì happy path chạy; phải đạt acceptance và regression liên quan.",
        "Bảo toàn `/api/v1` baseline trong lúc thêm `/api/v2`, chỉ remove sau E2E và quyết định compatibility.",
    ])

    add_h(doc, "18.2. Ordered backlog", 2)
    add_table(doc, ["ID", "Task", "Depends", "Done when"], [
        ["V2-001", "Contracts/models", "—", "JSON/Pydantic/API examples đồng nhất"],
        ["V2-002–003", "Employee/workspace context + assembler", "001", "Source/freshness đủ; no cross-case leak"],
        ["V2-004–005", "Intent + slots/confidence/clarification", "001–003", "Structured outputs; no-repeat target"],
        ["V2-006–007", "Ingestion/index + hybrid retrieval/matcher", "001,004", "Index reproducible; Hit@5/citation đạt gate"],
        ["V2-008", "Eligibility/Legal", "001,007", "Blocking/evidence đúng; unsafe pass 0"],
        ["V2-009", "Workflow/state/resume", "005,008", "DAG/retry/impact selection đúng"],
        ["V2-010", "Operations/dedup/artifacts", "009", "No duplicate artifacts/tasks"],
        ["V2-011", "Safety/approval/executor", "001,009,010", "No unsafe/duplicate write"],
        ["V2-012", "Storage/observability", "001,009,011", "Restart-safe pilot profile, sanitized trace"],
        ["V2-013", "API/UI", "002–012", "Complete ABC journey qua UI"],
        ["V2-014–015", "Evaluation + E2E hardening", "All", "Thresholds đo được; acceptance system pass"],
    ], widths=[1300, 3150, 2000, 2910], font_size=7.9)

    add_h(doc, "18.3. Vertical checkpoints", 2)
    add_bullets(doc, [
        "Checkpoint 1 — Understand only: “Kiểm tra còn thiếu gì” + workspace → đúng intent/slots, không hỏi customer/case.",
        "Checkpoint 2 — Grounded recommendation: ABC multi-intent → products + eligibility + citations.",
        "Checkpoint 3 — Controlled workflow: pending information → upload UBO → partial resume → approval → một mock action.",
        "Checkpoint 4 — Pilot-shaped app: persistent state, correction UI, trace và eval report.",
    ])

    add_h(doc, "18.4. Lộ trình triển khai theo giai đoạn", 2)
    add_table(doc, ["Giai đoạn", "Deliverables", "Gate ra"], [
        ["1. Contracts & Context", "Schema, shared state, IAM/workspace/CRM mock, context header", "Auto-fill ≥98%, leak=0"],
        ["2. Intent & Product RAG", "Taxonomy, extractor, clarification, ingestion/index/retrieval", "Intent/RAG gates MVP"],
        ["3. Eligibility & Workflow", "Rule registry, Legal RAG, DAG/state/resume", "Unsafe pass=0; resume ≥90%"],
        ["4. Operations & Approval", "Checklist/draft/dedup, token/executor", "Duplicate/unsafe action=0"],
        ["5. Persistence/UI/Eval", "PostgreSQL/vector DB, observability, UI E2E, datasets", "System acceptance pass"],
        ["6. Pilot", "5–10 RM, sandbox integrations, feedback/quality dashboard", "Pilot gates + governance sign-off"],
    ], widths=[2300, 4300, 2760], font_size=8.6)

    add_h(doc, "19. HIỆN TRẠNG, ĐÃ CÓ VÀ CHƯA CÓ", 1)
    add_p(doc, "Baseline repo hiện có là MVP FastAPI với deterministic workflow, synthetic data và demo UI. Contracts V2, Employee/Workspace Context và Context Assembler đã được triển khai; Intent V2 và các module sau chưa được nối vào runtime E2E. Tại thời điểm cập nhật, toàn bộ 73 automated tests đang pass. Vì vậy hướng an toàn cho ngày 18/07/2026 là harden vertical slice `/api/v1` đang chạy, không cố hoàn thành toàn bộ backlog V2 và không dùng từ “production-ready”.")
    add_table(doc, ["Hạng mục", "Baseline đã có", "V2 còn cần"], [
        ["Shared state", "Pydantic MVP", "JSON contract V2 + provenance + migration"],
        ["Planner", "Deterministic DAG", "Nối Context/Intent + impact resume"],
        ["Product retrieval", "Hash embedding, in-memory hybrid-lite", "PDF/Excel ingestion, persistent hybrid index, ACL/version"],
        ["Product module", "Deterministic MVP", "Intent contract, matcher/evidence versioning"],
        ["Legal", "Synthetic rules", "Rule registry + Legal RAG + effective dates"],
        ["Operations", "Checklist/email draft", "Artifact reuse, dedup, partial update"],
        ["Approval", "HMAC demo", "Payload hash, nonce, expiry, one-time use, RBAC"],
        ["API/UI", "FastAPI + demo UI", "Context endpoints, correction, intent/evidence/approval panels"],
        ["Storage/index", "In-memory/local demo", "PostgreSQL, vector DB, Redis, migrations"],
        ["Tests", "73 tests pass; gồm baseline + V2 contracts/context", "MVP golden data cho corporate sales + UI E2E + demo smoke test"],
    ], widths=[2350, 2900, 4110], font_size=8.4)

    add_h(doc, "20. RỦI RO, DỮ LIỆU CẦN CÓ VÀ QUYẾT ĐỊNH MỞ", 1)
    add_h(doc, "20.1. Rủi ro chính", 2)
    copy_source_table(doc, source, 31, "Bảng 11. Rủi ro và biện pháp kiểm soát.", widths=[2200, 3000, 4160], font_size=8.4)
    add_h(doc, "20.2. Dữ liệu cần có", 2)
    add_table(doc, ["Dataset", "MVP", "Pilot / owner"], [
        ["Product catalog/policies", "Synthetic 5–10 products/rules", "Catalog + policy version được Product/Risk ký nhận"],
        ["Legal/KYC/AML", "Synthetic 3–5 rules", "Current policies + Legal/Compliance owner"],
        ["SOP/SLA", "Synthetic templates", "Approved SOP/business calendar từ Operations"],
        ["Employee/IAM", "Mock roles/scopes", "SSO/IAM spec từ IT/Security"],
        ["CRM/DMS/task", "Synthetic companies/adapters", "Sandbox schema/API + idempotency/status query"],
        ["Conversations/eval", "Curated synthetic", "De-identified samples + dual-reviewed high-risk labels"],
    ], widths=[2500, 2850, 4010], font_size=8.6)

    add_h(doc, "20.3. Quyết định phải chốt trước pilot", 2)
    add_bullets(doc, [
        "Intent taxonomy nào phản ánh đúng công việc RM và ai là owner?",
        "Field nào bắt buộc ở từng workflow stage; validity window của KYC/BCTC/task result?",
        "Khi CRM và document mâu thuẫn, nguồn nào thắng?",
        "Action nào RM tự approve, action nào cần cấp khác; pending information có được gửi email sau RM approve không?",
        "Vector DB, embedding, model gateway và data egress policy chuẩn nội bộ?",
        "Tenant/branch isolation, retention, encryption, tamper-evident audit và on-call/SLO ownership?",
    ])
    add_callout(doc, "Cách AI coding xử lý điều chưa biết", "Dùng interface/mock và gắn nhãn ASSUMPTION hoặc DATA REQUIRED; không tự bịa endpoint, policy, SLA hay quyền. Mọi temporary default phải được ghi vào decision/deviation log và giữ adapter thay thế được.", fill=AMBER_LIGHT, border=AMBER)

    add_h(doc, "21. KẾT LUẬN VÀ PITCH", 1)
    add_p(doc, "Đề xuất V2 không thay đổi câu chuyện cốt lõi: SHB xây một đội chuyên gia AI đứng sau mỗi RM. Điểm trưởng thành của V2 là đội chuyên gia này không hoạt động như các chatbot độc lập. Họ cùng đọc một case state, làm việc theo contract, dùng đúng dữ liệu và tool, dừng đúng lúc, tái sử dụng công việc cũ và để RM giữ quyền quyết định cuối.")
    add_p(doc, "Khả năng tạo khác biệt lớn nhất không nằm ở việc LLM nói hay hơn, mà ở việc hệ thống hiểu context sát hơn: biết RM đang ở customer/case nào, đã làm gì, thiếu gì và bước tiếp theo là gì. Khi context được chuẩn hóa, intent chính xác hơn; khi intent có provenance và confidence, workflow ít hỏi lại hơn; khi workflow có evidence, dedup và approval, AI mới tạo được giá trị vận hành mà vẫn kiểm soát rủi ro.")
    add_callout(doc, "Pitch chốt", "Chúng tôi không xây chatbot cho khách hàng doanh nghiệp. Chúng tôi xây một Context-Aware Expert Workspace đứng sau mỗi RM: hiểu đúng công việc đang diễn ra, phối hợp Product–Legal–Operations, biến tri thức thành phương án và hành động có căn cứ, nhưng chỉ thực thi sau khi con người phê duyệt.", fill=TEAL_LIGHT, border=TEAL)

    add_h(doc, "PHỤ LỤC A — TOOL/API VÀ HỢP ĐỒNG TÁC VỤ", 1)
    copy_source_table(doc, source, 22, "Bảng A1. Tool/API tối thiểu theo domain.", widths=[1500, 3650, 4210], font_size=8)
    add_p(doc, "Trong V2, danh sách trên phải đi qua Tool Registry có caller allowlist, JSON schema, risk, approval_required, timeout/retry và idempotency policy. Tên tool là contract nghiệp vụ, production endpoint được triển khai bằng adapter sau khi có specification thật.")

    add_h(doc, "PHỤ LỤC B — NHIỆM VỤ XÂY DỰNG CHI TIẾT CHO CÁC KHỐI", 1)
    copy_source_table(doc, source, 16, "Bảng B1. Quy trình xây dựng chung cho mỗi module/agent.", widths=[1700, 7660], font_size=8.4)
    for idx, caption in [
        (17, "Bảng B2. Planner Agent – backlog chi tiết."),
        (18, "Bảng B3. Product Agent – backlog chi tiết."),
        (19, "Bảng B4. Legal Agent – backlog chi tiết."),
        (20, "Bảng B5. Operations Agent – backlog chi tiết."),
        (21, "Bảng B6. Evidence/Guardrail/HITL – backlog chi tiết."),
    ]:
        copy_source_table(doc, source, idx, caption, widths=[900, 2850, 5610], font_size=7.9)

    add_h(doc, "PHỤ LỤC C — DỮ LIỆU CHO TỪNG AGENT/MODULE", 1)
    for idx, caption, widths in [
        (9, "Bảng C1. Dữ liệu dùng chung.", [1900, 3700, 2500, 1260]),
        (10, "Bảng C2. Dữ liệu cho Planner.", [2300, 3900, 3160]),
        (11, "Bảng C3. Dữ liệu cho Product.", [2200, 4000, 3160]),
        (12, "Bảng C4. Dữ liệu cho Legal.", [2200, 4000, 3160]),
        (13, "Bảng C5. Dữ liệu cho Operations.", [2200, 4000, 3160]),
        (14, "Bảng C6. Dữ liệu cho Evidence/Guardrail.", [2200, 4000, 3160]),
    ]:
        copy_source_table(doc, source, idx, caption, widths=widths, font_size=7.9)

    add_h(doc, "PHỤ LỤC D — CATALOG SẢN PHẨM DOANH NGHIỆP MINH HỌA", 1)
    add_callout(doc, "Lưu ý dữ liệu", "Danh mục dưới đây được giữ lại từ proposal ban đầu để bảo toàn câu chuyện và độ chi tiết. Đây là taxonomy/minh họa cho thiết kế Product RAG, không xác nhận tên thương mại, điều kiện, biểu phí hay chính sách hiện hành của SHB. Trước pilot, Product/Risk phải map sang product_id, version, effective date và source chính thức.", fill=AMBER_LIGHT, border=AMBER)
    appendix_captions = {
        37: "Bảng D1. Giai đoạn phát triển của doanh nghiệp và nhóm nhu cầu.",
        38: "Bảng D2. Tài khoản và tiền gửi doanh nghiệp.",
        39: "Bảng D3. Ngân hàng số và quản trị giao dịch.",
        40: "Bảng D4. Cash Management.",
        41: "Bảng D5. Thu hộ và đối soát.",
        42: "Bảng D6. Chi hộ và thanh toán.",
        43: "Bảng D7. Payroll và dịch vụ nhân viên.",
        44: "Bảng D8. Vốn lưu động và tín dụng ngắn hạn.",
        45: "Bảng D9. Tín dụng đầu tư trung/dài hạn.",
        46: "Bảng D10. Bảo lãnh.",
        47: "Bảng D11. Thanh toán quốc tế và trade finance.",
        48: "Bảng D12. Ngoại hối.",
        49: "Bảng D13. Thẻ doanh nghiệp.",
        50: "Bảng D14. Supply Chain Finance.",
        51: "Bảng D15. Merchant acquiring và thanh toán bán hàng.",
    }
    for idx in range(37, 52):
        copy_source_table(
            doc, source, idx, appendix_captions[idx],
            widths=[4200, 5160], font_size=8.5, caption_before=True,
            keep_table_together=True,
        )

    add_h(doc, "PHỤ LỤC E — KẾ HOẠCH FAST-TRACK MVP TRONG 1 NGÀY", 1)
    add_p(doc, "Mục tiêu của ngày build không phải hoàn thành target architecture. Mục tiêu là tạo một câu chuyện bán hàng doanh nghiệp hoàn chỉnh, ổn định và có thể giải thích. Mỗi khối thời gian kết thúc bằng một checkpoint chạy được; nếu chậm, cắt P1 trước và giữ P0.")
    add_table(doc, ["Khung giờ", "P0 phải hoàn thành", "Artifact/acceptance"], [
        ["08:00–09:00", "Freeze case hero, 6 product IDs, 4 intents, 2 blocking documents và expected states", "scenario spec + data manifest; không đổi scope sau 09:00"],
        ["09:00–11:00", "Mở rộng synthetic catalog/rules/SOP; nối keyword/structured extraction tối thiểu", "Product/Legal/Ops outputs dùng cùng IDs; unit tests pass"],
        ["11:00–12:00", "Đảm bảo partial branch: transaction tiếp tục, credit pending_information", "E2E assertion cho branch status và missing UBO/BCTC"],
        ["13:00–15:00", "Thay raw JSON UI bằng cards: context, needs, bundle, eligibility, checklist, evidence, approval", "RM hiểu kết quả trong <2 phút; không cần đọc trace JSON"],
        ["15:00–16:00", "Tạo 10 golden scenarios và 5 security cases", "Regression runner pass; forbidden actions được assert"],
        ["16:00–17:00", "Demo rehearsal 3 lần từ clean start; kiểm tra restart và port", "Runbook có lệnh; thời lượng demo 5–7 phút"],
        ["17:00–18:00", "Fix blocker, chụp backup, đóng gói README và fallback video/screenshots nếu cần", "Release candidate; no critical known failure"],
    ], widths=[1600, 4700, 3060], font_size=8.2)
    add_h(doc, "Thứ tự cắt scope khi thiếu thời gian", 2)
    add_bullets(doc, [
        "Cắt trước: animation/UI polish, API Banking alternative, thêm customer thứ tư, LLM thật.",
        "Giữ bắt buộc: hero case, 6 products, branch blocking, evidence, RM approval, 10 golden cases và demo runbook.",
        "Không được cắt: nhãn SYNTHETIC DEMO DATA, guardrail external action, missing-information behavior và RM approval.",
    ])

    doc.add_page_break()
    add_h(doc, "PHỤ LỤC F — SYSTEM ACCEPTANCE SCENARIOS", 1)
    add_table(doc, ["ID", "Scenario", "Kết quả bắt buộc"], [
        ["AC-01", "Context-first short request", "Resolve customer/case/product từ workspace, không clarification, trả checklist"],
        ["AC-02", "ABC multi-intent", "Recommend catalog products; credit pending vì UBO/BCTC; giữ non-credit"],
        ["AC-03", "Partial resume", "Upload UBO chỉ rerun impacted nodes và update artifacts"],
        ["AC-04", "Deduplication", "Equivalent active task được reuse, không side effect thứ hai"],
        ["AC-05", "Approval integrity", "Edit làm token invalid; unchanged payload executes once despite retry"],
        ["AC-06", "Security", "Injection/tool privilege escalation bị block và audit high severity"],
    ], widths=[1100, 2800, 5460], font_size=8.6)

    configure_headers_footers(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
