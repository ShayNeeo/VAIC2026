"""File parser router for PDF, DOCX, XLSX, text and JSON ingestion."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from pydantic import BaseModel, ConfigDict


class ParsedSection(BaseModel):
    model_config = ConfigDict(extra="forbid")

    location: str
    text: str
    metadata: Dict[str, Any]


class UnsupportedDocumentError(ValueError):
    pass


def parse_document(path: str | Path) -> List[ParsedSection]:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(source)
    if suffix == ".docx":
        return _parse_docx(source)
    if suffix == ".xlsx":
        return _parse_xlsx(source)
    if suffix in {".txt", ".md", ".csv"}:
        return [ParsedSection(location="file", text=source.read_text(encoding="utf-8"), metadata={"type": suffix[1:]})]
    if suffix == ".json":
        payload = json.loads(source.read_text(encoding="utf-8"))
        return [ParsedSection(location="json", text=json.dumps(payload, ensure_ascii=False), metadata={"type": "json"})]
    raise UnsupportedDocumentError(f"unsupported document type: {suffix}")


def parse_document_bytes(filename: str, data: bytes) -> List[ParsedSection]:
    """Parse an upload without persisting the raw file to local disk."""

    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf_stream(BytesIO(data))
    if suffix == ".docx":
        return _parse_docx_stream(BytesIO(data))
    if suffix == ".xlsx":
        return _parse_xlsx_stream(BytesIO(data))
    if suffix in {".txt", ".md", ".csv"}:
        return [ParsedSection(location="file", text=data.decode("utf-8"), metadata={"type": suffix[1:]})]
    if suffix == ".json":
        payload = json.loads(data.decode("utf-8"))
        return [ParsedSection(location="json", text=json.dumps(payload, ensure_ascii=False), metadata={"type": "json"})]
    raise UnsupportedDocumentError(f"unsupported document type: {suffix}")


def ocr_pdf_sections(data: bytes, *, lang: str | None = None) -> List[ParsedSection]:
    """OCR fallback for a scanned/image PDF: one ParsedSection per page,
    same shape as _parse_pdf_stream() so callers (app/intake/service.py)
    can run extraction_quality() over the result identically. Raises
    app.knowledge.ocr.OcrUnavailableError if the OCR toolchain cannot run
    in this environment -- callers must not treat that the same as "OCR
    ran and found nothing"."""
    from app.knowledge.ocr import ocr_pdf_bytes

    pages = ocr_pdf_bytes(data, lang=lang)
    return [
        ParsedSection(
            location=f"page:{page['page']}", text=page["text"],
            metadata={"page": page["page"], "type": "pdf_ocr", "ocr_confidence": page["mean_confidence"]},
        )
        for page in pages
    ]


def extraction_quality(sections: List[ParsedSection]) -> Dict[str, Any]:
    non_empty = sum(bool(section.text.strip()) for section in sections)
    ratio = non_empty / len(sections) if sections else 0.0
    return {
        "section_count": len(sections),
        "non_empty_sections": non_empty,
        "non_empty_ratio": round(ratio, 4),
        "publishable": bool(sections) and ratio >= 0.8,
    }


def _parse_pdf(path: Path) -> List[ParsedSection]:
    from pypdf import PdfReader

    return _pdf_sections(PdfReader(str(path)))


def _parse_pdf_stream(stream: BytesIO) -> List[ParsedSection]:
    from pypdf import PdfReader

    return _pdf_sections(PdfReader(stream))


def _pdf_sections(reader) -> List[ParsedSection]:
    return [
        ParsedSection(
            location=f"page:{index}",
            text=page.extract_text() or "",
            metadata={"page": index, "type": "pdf"},
        )
        for index, page in enumerate(reader.pages, start=1)
    ]


def _parse_docx(path: Path) -> List[ParsedSection]:
    from docx import Document

    return _docx_sections(Document(str(path)))


def _parse_docx_stream(stream: BytesIO) -> List[ParsedSection]:
    from docx import Document

    return _docx_sections(Document(stream))


def _docx_sections(document) -> List[ParsedSection]:
    sections: List[ParsedSection] = []
    heading = "document"
    buffer: List[str] = []

    def flush() -> None:
        if buffer:
            sections.append(ParsedSection(location=heading, text="\n".join(buffer), metadata={"type": "docx"}))
            buffer.clear()

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            flush()
            heading = text
        else:
            buffer.append(text)
    flush()
    for table_index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        sections.append(
            ParsedSection(
                location=f"table:{table_index}", text="\n".join(rows),
                metadata={"type": "docx_table", "header_preserved": bool(rows)},
            )
        )
    return sections


def _parse_xlsx(path: Path) -> List[ParsedSection]:
    from openpyxl import load_workbook

    return _xlsx_sections(load_workbook(path, read_only=True, data_only=True))


def _parse_xlsx_stream(stream: BytesIO) -> List[ParsedSection]:
    from openpyxl import load_workbook

    return _xlsx_sections(load_workbook(stream, read_only=True, data_only=True))


def _xlsx_sections(workbook) -> List[ParsedSection]:
    sections: List[ParsedSection] = []
    try:
        for sheet in workbook.worksheets:
            rows = [" | ".join("" if value is None else str(value) for value in row) for row in sheet.iter_rows(values_only=True)]
            sections.append(
                ParsedSection(
                    location=f"sheet:{sheet.title}", text="\n".join(rows),
                    metadata={"type": "xlsx", "sheet": sheet.title, "header_preserved": bool(rows)},
                )
            )
    finally:
        workbook.close()
    return sections
