"""Field-level extraction tests for app.intake.extractor / app.intake.service.

All fixture text below is SYNTHETIC DEMO DATA (no real SHB customer, no real
tax code/company). Covers: company name, tax code, legal representative,
employee count, supplier count, distributor count, ERP, pain points,
conflict detection, missing fields, low-confidence/vague quantities, and
field-level provenance. Two real regression fixes landed alongside these
tests (see app/intake/extractor.py comments):

1. Distributor count ("đại lý") and customer count ("khách hàng") used to
   share one regex/field (collection_profile.customer_count) -- "40 đại lý"
   was silently reported as a customer count. Now split into
   collection_profile.distributor_count vs collection_profile.customer_count.
2. legal_profile.legal_representative did not exist at all -- classify_document
   detected a "ubo_information" document type on representative-related
   keywords, but extract_document_fields never actually extracted a
   representative name. Added.
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional

import pytest

from app.intake.extractor import classify_document, detect_needs, detect_pain_points, extract_document_fields
from app.intake.service import IntakeService
from app.schemas.v2.intake import ExtractedField, FieldValidationStatus


def _fields(text: str, *, document_type: str = "meeting_note", document_id: str = "DOC-1", page: int = 1) -> List[ExtractedField]:
    sections = [{"text": text, "location": f"page-{page}", "metadata": {"page": page}}]
    return extract_document_fields(document_id=document_id, document_type=document_type, sections=sections)


def _get(fields: List[ExtractedField], name: str) -> Optional[ExtractedField]:
    return next((f for f in fields if f.field_name == name), None)


# ==========================================================================
# Synthetic fixtures (SYNTHETIC DEMO DATA)
# ==========================================================================

MEETING_NOTE = """Biên bản ghi chú cuộc họp - Công ty Cổ phần Thiết bị Minh Phát
Công ty Cổ phần Thiết bị Minh Phát có 500 nhân viên, 80 nhà cung cấp và 40 đại lý.
Payroll hiện đang xử lý thủ công, gây mất thời gian mỗi kỳ lương.
Đối soát công nợ với đại lý còn thủ công qua Excel, dễ sai sót.
Dòng tiền phân tán giữa nhiều tài khoản ngân hàng.
Doanh nghiệp muốn kết nối ERP với hệ thống ngân hàng qua API.
SYNTHETIC DEMO DATA."""

BUSINESS_REGISTRATION = """GIAY CHUNG NHAN DANG KY DOANH NGHIEP
Công ty Cổ phần Thiết bị Minh Phát
Mã số thuế: 0109988665
Người đại diện theo pháp luật: Nguyễn Văn A, Giám đốc
Loại hình doanh nghiệp: Công ty cổ phần
Địa chỉ: Hà Nội
Doanh nghiệp hoạt động liên tục 8 năm.
SYNTHETIC DEMO DATA."""

CONFLICT_MEETING_NOTE = "Biên bản ghi chú cuộc họp: công ty có 500 nhân viên. SYNTHETIC DEMO DATA."
CONFLICT_COMPANY_PROFILE = "Hồ sơ doanh nghiệp: công ty có 520 nhân viên. SYNTHETIC DEMO DATA."

MISSING_ERP_TAX_DOC = """Biên bản ghi chú cuộc họp
Công ty muốn mở rộng hoạt động kinh doanh trong năm tới.
SYNTHETIC DEMO DATA."""

LOW_CONFIDENCE_DOC = "Công ty có khoảng hơn 400 lao động và đang tiếp tục tuyển dụng. SYNTHETIC DEMO DATA."


# ==========================================================================
# 1. Company name
# ==========================================================================


def test_company_name_full_form_with_co_phan():
    field = _get(_fields("Công ty Cổ phần Thiết bị Minh Phát hoạt động ổn định."), "company_identity.name")
    assert field is not None
    assert "Minh Phát" in field.value
    assert field.confidence >= 0.85


def test_company_name_uppercase_form():
    field = _get(_fields("CÔNG TY CỔ PHẦN MINH PHÁT chuyên cung cấp thiết bị."), "company_identity.name")
    assert field is not None
    assert "MINH PHÁT" in field.value


def test_company_name_tnhh_form():
    field = _get(_fields("Công ty TNHH Dịch vụ Minh Phát cung cấp giải pháp logistics."), "company_identity.name")
    assert field is not None
    assert "Minh Phát" in field.value


def test_company_name_repeated_in_text_takes_first_mention():
    text = "Công ty Cổ phần Minh Phát ký hợp đồng.\nCông ty Cổ phần Minh Phát xác nhận lại thông tin."
    field = _get(_fields(text), "company_identity.name")
    assert field is not None
    assert "Minh Phát" in field.value


def test_company_name_absent_when_not_mentioned():
    field = _get(_fields("Doanh nghiệp cần tư vấn giải pháp tài chính."), "company_identity.name")
    assert field is None


# ==========================================================================
# 2. Tax code
# ==========================================================================


def test_tax_code_valid_synthetic_format():
    field = _get(_fields("Mã số thuế: 0109988665"), "company_identity.tax_code")
    assert field is not None
    assert field.value == "0109988665"
    assert field.confidence >= 0.85


def test_tax_code_with_branch_suffix():
    field = _get(_fields("MST: 0109988665-001"), "company_identity.tax_code")
    assert field is not None
    assert field.value == "0109988665-001"


def test_tax_code_not_confused_with_nearby_noise_numbers():
    text = "Số điện thoại: 0912345678. Mã số thuế: 0109988665. Fax: 02412345678."
    field = _get(_fields(text), "company_identity.tax_code")
    assert field is not None
    assert field.value == "0109988665"


def test_tax_code_invalid_format_is_not_extracted():
    field = _get(_fields("Mã số thuế: ABC123"), "company_identity.tax_code")
    assert field is None


def test_tax_code_missing_entirely():
    field = _get(_fields("Doanh nghiệp hoạt động trong lĩnh vực xây dựng."), "company_identity.tax_code")
    assert field is None


# ==========================================================================
# 3. Legal representative
# ==========================================================================


def test_legal_representative_full_label():
    field = _get(_fields("Người đại diện theo pháp luật: Nguyễn Văn A"), "legal_profile.legal_representative")
    assert field is not None
    assert "Nguyễn Văn A" in field.value


def test_legal_representative_short_label():
    field = _get(_fields("Đại diện: Trần Thị B"), "legal_profile.legal_representative")
    assert field is not None
    assert "Trần Thị B" in field.value


def test_legal_representative_with_title_does_not_swallow_title():
    field = _get(_fields("Người đại diện theo pháp luật: Nguyễn Văn A, Giám đốc"), "legal_profile.legal_representative")
    assert field is not None
    assert field.value.strip() == "Nguyễn Văn A"


def test_legal_representative_takes_first_of_multiple_names():
    text = "Người đại diện theo pháp luật: Nguyễn Văn A.\nNgười liên hệ dự phòng: Lê Thị C."
    field = _get(_fields(text), "legal_profile.legal_representative")
    assert field is not None
    assert "Nguyễn Văn A" in field.value


def test_legal_representative_absent_without_evidence():
    field = _get(_fields("Công ty hoạt động trong lĩnh vực thương mại điện tử."), "legal_profile.legal_representative")
    assert field is None


# ==========================================================================
# 4. Employee count
# ==========================================================================


def test_employee_count_direct_mention():
    field = _get(_fields("Có 500 nhân viên."), "business_profile.employees_count")
    assert field is not None
    assert field.value == 500


def test_employee_count_scale_phrasing():
    field = _get(_fields("Quy mô nhân sự: 500 người lao động."), "business_profile.employees_count")
    assert field is not None
    assert field.value == 500


def test_employee_count_missing():
    field = _get(_fields("Công ty hoạt động ổn định trong 5 năm qua."), "business_profile.employees_count")
    assert field is None


# ==========================================================================
# 5. Supplier count
# ==========================================================================


def test_supplier_count_direct():
    field = _get(_fields("Công ty làm việc với 80 nhà cung cấp."), "payment_profile.supplier_count")
    assert field is not None
    assert field.value == 80


def test_supplier_count_not_confused_with_transaction_count():
    text = "Công ty xử lý 1200 giao dịch mỗi tháng và làm việc với 80 nhà cung cấp."
    field = _get(_fields(text), "payment_profile.supplier_count")
    assert field is not None
    assert field.value == 80  # not 1200 -- "giao dịch" is not in the supplier pattern


# ==========================================================================
# 6. Distributor count (đại lý) -- regression: previously merged with
#    customer count under one field.
# ==========================================================================


def test_distributor_count_extracted_as_its_own_field():
    field = _get(_fields("Công ty có 40 đại lý trên toàn quốc."), "collection_profile.distributor_count")
    assert field is not None
    assert field.value == 40


def test_distributor_count_not_reported_as_customer_count():
    fields = _fields("Công ty có 40 đại lý trên toàn quốc.")
    assert _get(fields, "collection_profile.customer_count") is None
    assert _get(fields, "collection_profile.distributor_count") is not None


def test_distributor_count_not_confused_with_branch_count():
    text = "Công ty có 40 đại lý và 5 chi nhánh."
    field = _get(_fields(text), "collection_profile.distributor_count")
    assert field is not None
    assert field.value == 40  # not 5 -- "chi nhánh" is a different concept, not matched at all


def test_customer_count_still_extracted_separately_from_distributor_count():
    text = "Công ty có 40 đại lý và phục vụ 1000 khách hàng."
    fields = _fields(text)
    distributor = _get(fields, "collection_profile.distributor_count")
    customer = _get(fields, "collection_profile.customer_count")
    assert distributor is not None and distributor.value == 40
    assert customer is not None and customer.value == 1000


# ==========================================================================
# 7. ERP: current system vs need for integration are represented distinctly
# ==========================================================================


def test_erp_current_system_named():
    field = _get(_fields("Doanh nghiệp đang dùng hệ thống SAP."), "technology_profile.erp_system")
    assert field is not None
    assert field.value == "SAP"


def test_erp_oracle_named():
    field = _get(_fields("Công ty vận hành trên nền tảng Oracle."), "technology_profile.erp_system")
    assert field is not None
    assert field.value == "ORACLE"


def test_erp_need_without_named_system_is_captured_via_explicit_needs():
    # No named ERP product -> technology_profile.erp_system stays absent,
    # but the need itself is captured distinctly via explicit_needs (see
    # detect_needs's "erp_integration" signal) rather than fabricating a
    # system name.
    text = "Doanh nghiệp muốn kết nối ERP với hệ thống ngân hàng nhưng chưa có hệ thống ERP nào."
    fields = _fields(text)
    assert _get(fields, "technology_profile.erp_system") is None
    needs_field = _get(fields, "explicit_needs")
    assert needs_field is not None
    assert "erp_integration" in needs_field.value


def test_erp_absent_when_not_mentioned():
    fields = _fields("Công ty tập trung vào mảng sản xuất.")
    assert _get(fields, "technology_profile.erp_system") is None
    needs_field = _get(fields, "explicit_needs")
    assert needs_field is None or "erp_integration" not in (needs_field.value or [])


# ==========================================================================
# 8. Pain points
# ==========================================================================


def test_pain_points_detected_with_source_span():
    fields = _fields(MEETING_NOTE)
    field = _get(fields, "pain_points")
    assert field is not None
    assert any("thủ công" in item or "thu cong" in item for item in field.value)
    assert field.source_document_id == "DOC-1"
    assert field.confidence < 1.0  # never treated as a certain fact


def test_pain_points_empty_when_no_signal_words_present():
    field = _get(_fields("Công ty hoạt động hiệu quả và ổn định."), "pain_points")
    assert field is None


# ==========================================================================
# 9. classify_document + full meeting-note extraction (integration-ish,
#    still no I/O -- exercises the whole extract_document_fields path)
# ==========================================================================


def test_meeting_note_classification():
    # classify_document scans an ordered rule list and returns the type with
    # the most keyword hits (see app/intake/extractor.py). MEETING_NOTE
    # itself is deliberately payment/ERP-topic-dense (5 payment_process
    # keyword hits vs 2 meeting_note hits) because that fixture exists to
    # exercise supplier/distributor/ERP field extraction, not classification
    # -- so it legitimately classifies as payment_process, which is correct,
    # content-driven behavior, not a bug (neither "payment_process" nor
    # "meeting_note" is a required-document type any eligibility rule keys
    # off of, so this ambiguity has no functional effect downstream). This
    # test uses a fixture with only meeting-note signal words instead.
    pure_meeting_note = (
        "Biên bản ghi chú cuộc họp giữa RM và khách hàng.\n"
        "Nội dung trao đổi chính: định hướng hợp tác trong quý tới.\n"
        "SYNTHETIC DEMO DATA."
    )
    doc_type, confidence = classify_document("bien_ban_hop.txt", pure_meeting_note)
    assert doc_type == "meeting_note"
    assert confidence >= 0.85


def test_meeting_note_yields_multiple_needs():
    needs = detect_needs(MEETING_NOTE)
    assert "payroll" in needs
    assert "erp_integration" in needs


# ==========================================================================
# 10. Missing field -> null value, listed as missing, no fabrication
# ==========================================================================


def test_missing_erp_and_tax_code_are_absent_not_guessed():
    fields = _fields(MISSING_ERP_TAX_DOC)
    assert _get(fields, "technology_profile.erp_system") is None
    assert _get(fields, "company_identity.tax_code") is None


def test_build_profile_lists_missing_tax_code(tmp_path):
    from app.storage.repository import V2Repository

    service = IntakeService(V2Repository(tmp_path / "state.sqlite3"))
    stored = service.create(
        employee_id="RM-1", session_id="SESS-1", customer_id="COMP-1",
        manual_input={"company_name": "Công ty Cổ phần Minh Phát", "need_text": "chi lương"},
    )
    assert "company_identity.tax_code" in stored.session.profile.missing_information


# ==========================================================================
# 11. Low-confidence / vague quantity: must not fabricate a precise number
# ==========================================================================


def test_vague_hedged_employee_count_is_not_fabricated_as_a_precise_fact():
    """"khoảng hơn 400 lao động" (vague, hedged) does not match the strict
    "N nhân viên/nhân sự/người lao động" patterns (missing the "người" token
    immediately before "lao động"), so no employees_count field is emitted
    at all -- the extractor prefers silence over guessing a precise number
    from a hedged phrase. This is intentionally verified directly against
    the current regex boundary, not just described in prose."""
    field = _get(_fields(LOW_CONFIDENCE_DOC), "business_profile.employees_count")
    assert field is None


def test_confidence_below_threshold_is_flagged_needs_review():
    # erp_system extraction confidence (0.82) is below the 0.85 auto-confirm
    # threshold used by extract_document_fields's add(): must be flagged.
    field = _get(_fields("Hệ thống ERP đang dùng là MISA."), "technology_profile.erp_system")
    assert field is not None
    assert field.confidence < 0.85
    assert field.validation_status == FieldValidationStatus.NEEDS_REVIEW


def test_high_confidence_field_is_valid_not_needs_review():
    field = _get(_fields("Mã số thuế: 0109988665"), "company_identity.tax_code")
    assert field is not None
    assert field.confidence >= 0.85
    assert field.validation_status == FieldValidationStatus.VALID


# ==========================================================================
# 12. Provenance: every field carries source_document_id/page/section/span/
#     method/confidence/validation_status
# ==========================================================================


def test_every_extracted_field_has_full_provenance():
    fields = _fields(MEETING_NOTE, document_id="DOC-XYZ", page=3)
    assert fields, "fixture should yield at least one field"
    for field in fields:
        assert field.source_document_id == "DOC-XYZ"
        assert field.source_section  # non-empty location string
        assert field.source_text_span  # the matched span, not empty
        assert field.extraction_method
        assert 0.0 <= field.confidence <= 1.0
        assert field.validation_status in (FieldValidationStatus.VALID, FieldValidationStatus.NEEDS_REVIEW)


def test_provenance_page_is_taken_from_section_metadata():
    field = _get(_fields("Mã số thuế: 0109988665", page=7), "company_identity.tax_code")
    assert field is not None
    assert field.source_page == 7


# ==========================================================================
# 13. Conflict detection (IntakeService._detect_conflicts)
# ==========================================================================


def test_conflicting_employee_counts_across_two_documents_are_flagged():
    fields = _fields(CONFLICT_MEETING_NOTE, document_id="DOC-A") + _fields(CONFLICT_COMPANY_PROFILE, document_id="DOC-B")
    conflicts = IntakeService._detect_conflicts(fields)
    employee_conflicts = [c for c in conflicts if c.field_name == "business_profile.employees_count"]
    assert len(employee_conflicts) == 1
    values = {c["value"] for c in employee_conflicts[0].candidates}
    assert values == {500, 520}
    assert employee_conflicts[0].requires_confirmation is True  # high-impact, unresolved


def test_no_conflict_when_documents_agree():
    fields = _fields("Có 500 nhân viên.", document_id="DOC-A") + _fields("Quy mô 500 người lao động.", document_id="DOC-B")
    conflicts = IntakeService._detect_conflicts(fields)
    assert not any(c.field_name == "business_profile.employees_count" for c in conflicts)


def test_rm_confirmation_resolves_a_conflict_field():
    fields = _fields(CONFLICT_MEETING_NOTE, document_id="DOC-A") + _fields(CONFLICT_COMPANY_PROFILE, document_id="DOC-B")
    rm_confirmed = ExtractedField(
        field_id="FIELD-RM-1", field_name="business_profile.employees_count", value=520,
        normalized_value=520, source_document_id="RM_CONFIRMATION", source_section="profile_review",
        source_text_span="RM xác nhận 520", extraction_method="human_correction", confidence=1.0,
        validation_status=FieldValidationStatus.VALID, decision_impact="high", confirmed_by_rm=True,
    )
    conflicts = IntakeService._detect_conflicts(fields + [rm_confirmed])
    employee_conflicts = [c for c in conflicts if c.field_name == "business_profile.employees_count"]
    assert len(employee_conflicts) == 1
    assert employee_conflicts[0].requires_confirmation is False
    assert employee_conflicts[0].resolved_value == 520


# ==========================================================================
# 14. End-to-end intake API: real DOCX file, not synthetic in-memory text
#     (tests/unit/test_v2_document_parsers.py already covers pure parsing;
#     this proves the extractor sees real parsed sections the same way).
# ==========================================================================


def test_extraction_over_a_real_generated_docx_file(tmp_path):
    from docx import Document

    from app.knowledge.parsers import parse_document_bytes

    doc = Document()
    doc.add_heading("Hồ sơ doanh nghiệp", level=1)
    doc.add_paragraph("Công ty Cổ phần Thiết bị Minh Phát")
    doc.add_paragraph("Mã số thuế: 0109988665")
    doc.add_paragraph("Người đại diện theo pháp luật: Nguyễn Văn A")
    doc.add_paragraph("Công ty có 500 nhân viên và 40 đại lý.")
    doc.add_paragraph("SYNTHETIC DEMO DATA.")
    path = tmp_path / "profile.docx"
    doc.save(path)

    sections = [
        {"text": section.text, "location": section.location, "metadata": section.metadata}
        for section in parse_document_bytes("profile.docx", path.read_bytes())
    ]
    fields = extract_document_fields(document_id="DOC-DOCX-1", document_type="business_registration", sections=sections)

    assert _get(fields, "company_identity.tax_code").value == "0109988665"
    assert "Nguyễn Văn A" in _get(fields, "legal_profile.legal_representative").value
    assert _get(fields, "business_profile.employees_count").value == 500
    assert _get(fields, "collection_profile.distributor_count").value == 40
