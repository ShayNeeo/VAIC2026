"""Parser routing and extraction-quality tests."""

from __future__ import annotations

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from app.knowledge.parsers import extraction_quality, parse_document, parse_document_bytes


def test_docx_preserves_heading_and_table_header(tmp_path):
    path = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("Điều kiện Payroll", level=1)
    document.add_paragraph("Doanh nghiệp có tối thiểu 10 nhân sự.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Trường"
    table.rows[0].cells[1].text = "Giá trị"
    table.rows[1].cells[0].text = "Nhân sự"
    table.rows[1].cells[1].text = "10"
    document.save(path)
    sections = parse_document(path)
    assert sections[0].location == "Điều kiện Payroll"
    assert "10 nhân sự" in sections[0].text
    assert sections[1].metadata["header_preserved"] is True


def test_xlsx_preserves_sheet_and_units(tmp_path):
    path = tmp_path / "fees.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Fees"
    sheet.append(["product_id", "amount", "currency"])
    sheet.append(["PROD-PAYROLL", 0, "VND"])
    workbook.save(path)
    sections = parse_document(path)
    assert sections[0].location == "sheet:Fees"
    assert "amount | currency" in sections[0].text
    assert "VND" in sections[0].text


def test_blank_pdf_is_parsed_but_fails_quality_gate(tmp_path):
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as handle:
        writer.write(handle)
    sections = parse_document(path)
    report = extraction_quality(sections)
    assert report["section_count"] == 1
    assert report["publishable"] is False


def test_upload_bytes_are_parsed_without_a_persistent_file():
    sections = parse_document_bytes("policy.txt", "Điều kiện UBO".encode("utf-8"))
    assert sections[0].text == "Điều kiện UBO"
